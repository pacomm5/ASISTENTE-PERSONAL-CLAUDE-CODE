#!/usr/bin/env /opt/homebrew/bin/python3
"""
Genera informes de seguimiento Word (.docx) — Abril 2026.
ICC, CX y Desgaste Instalaciones pendientes: esas secciones quedan como
"Sin datos registrados" hasta que lleguen los rankings.
"""

import os, datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

PERIODO   = "2026-04"
OUTPUT    = "/Users/franciscomartin/Documents/ASISTENTE PERSONAL DE CLAUDE CODE/informes"
HOY       = datetime.date.today().strftime("%d/%m/%Y")

# ══════════════════════════════════════════════════════════════════════════════
# DATOS
# ══════════════════════════════════════════════════════════════════════════════

# Medias desgaste asesores (18 activos, excl. JORGE)
MCAT = {"Amortiguadores":19.6,"Baterías":11.0,"Discos":20.6,"Escobillas":52.9,
        "Neumáticos":75.7,"Parabrisas":17.1,"Pastillas":49.5}
MTOTAL = 246.6
MFLD   = 11.1
MFID   = 63.48

# Desgaste asesores — (pos_18, total, {cat: val})
DAS = {
    "alvaro":            (7,  266.09, {"Amortiguadores":42.61,"Baterías":10.43,"Discos":24.35,"Escobillas":47.39,"Neumáticos":67.83,"Parabrisas":12.61,"Pastillas":60.87}),
    "guillermo":         (2,  323.46, {"Amortiguadores":35.80,"Baterías":16.05,"Discos":12.35,"Escobillas":76.54,"Neumáticos":93.83,"Parabrisas":30.86,"Pastillas":58.02}),
    "jose_maria_vazquez":(4,  283.43, {"Amortiguadores":27.43,"Baterías":9.14, "Discos":27.43,"Escobillas":63.43,"Neumáticos":100.57,"Parabrisas":5.14,"Pastillas":50.29}),
    "estefania":         (1,  341.67, {"Amortiguadores":8.33, "Baterías":25.00,"Discos":33.33,"Escobillas":8.33, "Neumáticos":108.33,"Parabrisas":125.00,"Pastillas":33.33}),
    "magan":             (3,  309.38, {"Amortiguadores":31.25,"Baterías":18.75,"Discos":32.81,"Escobillas":48.44,"Neumáticos":95.31,"Parabrisas":14.06,"Pastillas":68.75}),
    "javier_diaz":       (10, 247.54, {"Amortiguadores":15.57,"Baterías":9.84, "Discos":15.57,"Escobillas":54.10,"Neumáticos":96.72,"Parabrisas":11.48,"Pastillas":44.26}),
    "alberto_martinez":  (5,  272.20, {"Amortiguadores":18.26,"Baterías":11.20,"Discos":17.43,"Escobillas":80.08,"Neumáticos":92.12,"Parabrisas":2.49, "Pastillas":50.62}),
    "juan":              (8,  254.62, {"Amortiguadores":31.93,"Baterías":8.40, "Discos":24.79,"Escobillas":73.11,"Neumáticos":68.91,"Parabrisas":7.98, "Pastillas":39.50}),
    "taboada":           (11, 244.44, {"Amortiguadores":7.94, "Baterías":15.87,"Discos":26.98,"Escobillas":55.56,"Neumáticos":68.25,"Parabrisas":21.43,"Pastillas":48.41}),
    "alejandro":         (12, 234.44, {"Amortiguadores":11.11,"Baterías":13.33,"Discos":11.11,"Escobillas":70.00,"Neumáticos":70.00,"Parabrisas":17.78,"Pastillas":41.11}),
    "codru":             (9,  250.75, {"Amortiguadores":32.84,"Baterías":11.44,"Discos":21.89,"Escobillas":56.22,"Neumáticos":62.19,"Parabrisas":13.43,"Pastillas":52.74}),
    "carlos_aguilar":    (15, 202.46, {"Amortiguadores":0.82, "Baterías":4.51, "Discos":7.38, "Escobillas":77.05,"Neumáticos":64.34,"Parabrisas":4.10, "Pastillas":44.26}),
    "dennis":            (6,  267.68, {"Amortiguadores":39.39,"Baterías":6.06, "Discos":36.36,"Escobillas":62.63,"Neumáticos":51.52,"Parabrisas":7.07, "Pastillas":64.65}),
    "javier_pina":       (16, 200.77, {"Amortiguadores":7.69, "Baterías":16.15,"Discos":20.00,"Escobillas":24.62,"Neumáticos":71.54,"Parabrisas":18.46,"Pastillas":42.31}),
    "raul_munoz":        (14, 204.97, {"Amortiguadores":11.05,"Baterías":9.94, "Discos":16.57,"Escobillas":54.70,"Neumáticos":59.67,"Parabrisas":1.10, "Pastillas":51.93}),
    "jose_maria":        (13, 212.56, {"Amortiguadores":16.08,"Baterías":6.03, "Discos":17.09,"Escobillas":11.56,"Neumáticos":92.96,"Parabrisas":5.03, "Pastillas":63.82}),
    "jon":               (17, 191.49, {"Amortiguadores":6.38, "Baterías":3.19, "Discos":17.02,"Escobillas":50.00,"Neumáticos":67.02,"Parabrisas":7.98, "Pastillas":39.89}),
    "nuria":             (18, 130.32, {"Amortiguadores":9.03, "Baterías":3.23, "Discos":9.03, "Escobillas":38.71,"Neumáticos":32.26,"Parabrisas":1.29, "Pastillas":36.77}),
}

# Iron Man asesores — {fld, cal1, fid (None=N/A), horch (None=N/A)}
IAS = {
    "alvaro":            {"fld":50,  "cal1":5.00,"fid":61.65,"horch":25},
    "guillermo":         {"fld":15,  "cal1":6.00,"fid":60.53,"horch":102},
    "jose_maria_vazquez":{"fld":15,  "cal1":6.00,"fid":64.78,"horch":32},
    "estefania":         {"fld":0,   "cal1":5.00,"fid":None, "horch":None},
    "magan":             {"fld":9,   "cal1":6.00,"fid":63.53,"horch":51},
    "javier_diaz":       {"fld":9,   "cal1":6.00,"fid":66.98,"horch":32},
    "alberto_martinez":  {"fld":6,   "cal1":7.00,"fid":69.54,"horch":250},
    "juan":              {"fld":25,  "cal1":6.00,"fid":None, "horch":32},
    "taboada":           {"fld":11,  "cal1":6.00,"fid":59.65,"horch":65},
    "alejandro":         {"fld":5,   "cal1":6.00,"fid":65.79,"horch":60},
    "codru":             {"fld":4,   "cal1":5.00,"fid":65.12,"horch":145},
    "carlos_aguilar":    {"fld":6,   "cal1":7.00,"fid":65.74,"horch":174},
    "dennis":            {"fld":10,  "cal1":5.00,"fid":None, "horch":158},
    "javier_pina":       {"fld":6,   "cal1":6.00,"fid":57.98,"horch":102},
    "raul_munoz":        {"fld":5,   "cal1":5.00,"fid":57.84,"horch":107},
    "jose_maria":        {"fld":7,   "cal1":5.00,"fid":66.07,"horch":244},
    "jon":               {"fld":15,  "cal1":6.00,"fid":None, "horch":163},
    "nuria":             {"fld":1,   "cal1":6.00,"fid":None, "horch":102},
}

# ONE asesores — score numérico, cam/dif/multi como string con %
ONE_AM = {"score":6.2,"cam":8,"dif":16,"multi":6.1}
OAS = {
    "alberto_martinez":  {"score":6.3,"cam":"0%", "dif":"2%", "multi":"5,3"},
    "alejandro":         {"score":3.2,"cam":"0%", "dif":"10%","multi":"5,1"},
    "codru":             {"score":6.5,"cam":"7%", "dif":"15%","multi":"7,9"},
    "alvaro":            {"score":7.2,"cam":"22%","dif":"42%","multi":"6,6"},
    "carlos_aguilar":    {"score":7.5,"cam":"33%","dif":"55%","multi":"6,8"},
    "dennis":            {"score":5.3,"cam":"5%", "dif":"30%","multi":"6,8"},
    "estefania":         {"score":4.0,"cam":"13%","dif":"20%","multi":"4,0"},
    "taboada":           {"score":3.9,"cam":"0%", "dif":"11%","multi":"1,2"},
    "javier_pina":       {"score":4.5,"cam":"11%","dif":"1%", "multi":"5,0"},
    "guillermo":         {"score":6.0,"cam":"5%", "dif":"2%", "multi":"7,2"},
    "javier_diaz":       {"score":5.6,"cam":"1%", "dif":"2%", "multi":"5,8"},
    "magan":             {"score":5.5,"cam":"3%", "dif":"11%","multi":"6,0"},
    "jon":               {"score":6.9,"cam":"17%","dif":"10%","multi":"11,1"},
    "jose_maria":        {"score":7.0,"cam":"13%","dif":"24%","multi":"5,0"},
    "jose_maria_vazquez":{"score":7.7,"cam":"10%","dif":"15%","multi":"6,0"},
    "juan":              {"score":7.2,"cam":"2%", "dif":"13%","multi":"6,4"},
    "nuria":             {"score":8.7,"cam":"5%", "dif":"8%", "multi":"8,2"},
    "raul_munoz":        {"score":7.9,"cam":"4%", "dif":"15%","multi":"5,1"},
}

# ONE instalaciones
ONE_IM = {"score":7.3,"cam":8,"dif":16,"multi":6.2}
OINST = {
    "rivas_audi":  {"label":"Rivas Audi (00158)",         "score":6.8,"cam":"13%","dif":"30%","multi":"6,8"},
    "ayala":       {"label":"Ayala (00159)",               "score":9.2,"cam":"15%","dif":"25%","multi":"6,0"},
    "canarias":    {"label":"Rivas Canarias (00160)",      "score":5.6,"cam":"2%", "dif":"10%","multi":"5,3"},
    "vara":        {"label":"Vara (00982)",                "score":7.3,"cam":"5%", "dif":"10%","multi":"6,8"},
    "rivas_vw":    {"label":"Rivas VW / Industriales (30070)","score":6.1,"cam":"7%","dif":"3%","multi":"6,6"},
    "moncloa":     {"label":"Moncloa (31523)",             "score":9.0,"cam":"8%", "dif":"19%","multi":"5,8"},
}

# Desgaste instalaciones — (pos_7, total, {cat: val})
MCAT_INST = {"Amortiguadores":17.96,"Baterías":9.45,"Discos":18.44,"Escobillas":50.78,
             "Neumáticos":72.66,"Parabrisas":7.61,"Pastillas":50.60}
MTOTAL_INST = 227.50
DIT = {
    "emilio_rivas_audi":       (1, 264.34, {"Amortiguadores":38.17,"Baterías":10.75,"Discos":25.45,"Escobillas":51.97,"Neumáticos":62.90,"Parabrisas":17.74,"Pastillas":57.35}),
    "carlos_vara":             (2, 262.77, {"Amortiguadores":23.99,"Baterías":9.50, "Discos":22.74,"Escobillas":67.60,"Neumáticos":86.45,"Parabrisas":8.41, "Pastillas":44.08}),
    "izquierdo_canarias":      (3, 241.11, {"Amortiguadores":14.23,"Baterías":11.86,"Discos":22.13,"Escobillas":54.55,"Neumáticos":73.91,"Parabrisas":15.42,"Pastillas":49.01}),
    "alberto_martinez_ayala":  (4, 233.40, {"Amortiguadores":9.48, "Baterías":7.84, "Discos":12.37,"Escobillas":78.56,"Neumáticos":77.73,"Parabrisas":0.00, "Pastillas":47.42}),
    "jose_maria_moncloa_vw":   (5, 211.46, {"Amortiguadores":14.58,"Baterías":9.90, "Discos":17.71,"Escobillas":48.44,"Neumáticos":64.58,"Parabrisas":0.00, "Pastillas":56.25}),
    "jose_maria_moncloa_audi": (6, 199.47, {"Amortiguadores":12.77,"Baterías":5.85, "Discos":15.96,"Escobillas":15.43,"Neumáticos":89.36,"Parabrisas":0.00, "Pastillas":60.11}),
    "fernando_rivas_vw":       (7, 179.92, {"Amortiguadores":12.50,"Baterías":10.45,"Discos":12.70,"Escobillas":38.93,"Neumáticos":53.69,"Parabrisas":11.68,"Pastillas":39.96}),
}

# Iron Man jefes taller
IJT = {
    "emilio_rivas_audi":       {"icc_pos":3,  "cal1":5.00,"prod":92.45, "obj_prod":93, "var_rec":32.42,"var_mo":34.42,"horch":38},
    "alberto_martinez_ayala":  {"icc_pos":6,  "cal1":7.00,"prod":103.11,"obj_prod":105,"var_rec":0.17, "var_mo":21.40,"horch":39},
    "carlos_vara":             {"icc_pos":24, "cal1":6.00,"prod":91.23, "obj_prod":93, "var_rec":26.48,"var_mo":6.39, "horch":32},
    "izquierdo_canarias":      {"icc_pos":1,  "cal1":6.00,"prod":93.94, "obj_prod":93, "var_rec":7.48, "var_mo":3.61, "horch":46},
    "jose_maria_moncloa_audi": {"icc_pos":12, "cal1":4.00,"prod":90.15, "obj_prod":93, "var_rec":24.06,"var_mo":54.97,"horch":111},
    "jose_maria_moncloa_vw":   {"icc_pos":60, "cal1":5.00,"prod":90.15, "obj_prod":93, "var_rec":42.44,"var_mo":33.20,"horch":102},
    "fernando_rivas_vw":       {"icc_pos":14, "cal1":6.00,"prod":96.41, "obj_prod":89, "var_rec":3.03, "var_mo":-7.24,"horch":107},
    "luis_ramos_rivas_a_chapa":{"icc_pos":15, "cal1":5.00,"prod":117.15,"obj_prod":120,"var_rec":41.80,"var_mo":40.64,"horch":38},
    "luis_ramos_rivas_w_chapa":{"icc_pos":23, "cal1":6.00,"prod":117.15,"obj_prod":120,"var_rec":12.16,"var_mo":18.52,"horch":107},
    "pericles_canarias_chapa": {"icc_pos":16, "cal1":6.00,"prod":119.70,"obj_prod":120,"var_rec":-11.95,"var_mo":-11.70,"horch":46},
    "pericles_vara_chapa":     {"icc_pos":92, "cal1":6.00,"prod":119.70,"obj_prod":120,"var_rec":-28.45,"var_mo":-5.28,"horch":32},
}

# CX Instalaciones Abril 2026
_AU_REF = {
    "CEM Q4":       "0pts <4,60 | 1pt ≥4,60 | 2pts ≥4,80",
    "EC28":         "0pts >5% | 1pt ≤5% | 2pts ≤1%",
    "TRS":          "0pts <80% | 1pt ≥80% | 2pts ≥90%",
    "Conectividad": "0pts <25% | 1pt ≥25% | 2pts ≥45%",
}
_VW_REF = {
    "GBP":          "0pts <80% | 1pt ≥80% | 2pts 100%",
    "EC28":         "0pts >5% | 1pt ≤5% | 2pts ≤3%",
    "Conectividad": "0pts <25% | 1pt ≥25% | 2pts ≥40%",
    "Club VW":      "0pts <3% | 1pt ≥3% | 2pts ≥6%",
}
_LCV_REF = {
    "Sat. Gral":        "0pts <4,52 | 1pt ≥4,52 | 2pts ≥4,62",
    "EC28":             "0pts >5% | 1pt ≤5% | 2pts ≤3%",
    "Conectividad":     "0pts <10% | 1pt ≥10% | 2pts ≥25%",
    "Servicios Rápidos":"0pts <40% | 1pt ≥40% | 2pts ≥45%",
}

CX_APR = {
    "51AQ1 AU": {"label":"Rivas Audi (51AQ1 AU) — Audi Service", "total":5, "kpis":[
        ("Satisfacción General (CEM Q4)", "4,69",    1, _AU_REF["CEM Q4"]),
        ("Medidas de Servicio (EC28)",    "0,00%",   2, _AU_REF["EC28"]),
        ("Información Proceso (TRS)",     "62,50%",  0, _AU_REF["TRS"]),
        ("Conectividad Posventa",         "47,76%",  2, _AU_REF["Conectividad"]),
    ]},
    "51AQ2 AU": {"label":"Ayala (51AQ2 AU) — Audi Service", "total":7, "kpis":[
        ("Satisfacción General (CEM Q4)", "4,91",    2, _AU_REF["CEM Q4"]),
        ("Medidas de Servicio (EC28)",    "0,00%",   2, _AU_REF["EC28"]),
        ("Información Proceso (TRS)",     "82,61%",  1, _AU_REF["TRS"]),
        ("Conectividad Posventa",         "47,30%",  2, _AU_REF["Conectividad"]),
    ]},
    "51AQ3 AU": {"label":"Canarias (51AQ3 AU) — Audi Service", "total":5, "kpis":[
        ("Satisfacción General (CEM Q4)", "4,80",    2, _AU_REF["CEM Q4"]),
        ("Medidas de Servicio (EC28)",    "1,72%",   1, _AU_REF["EC28"]),
        ("Información Proceso (TRS)",     "87,50%",  1, _AU_REF["TRS"]),
        ("Conectividad Posventa",         "38,71%",  1, _AU_REF["Conectividad"]),
    ]},
    "31523 AU": {"label":"Moncloa Audi (31523 AU) — Audi Service", "total":4, "kpis":[
        ("Satisfacción General (CEM Q4)", "4,77",    1, _AU_REF["CEM Q4"]),
        ("Medidas de Servicio (EC28)",    "0,00%",   2, _AU_REF["EC28"]),
        ("Información Proceso (TRS)",     "60,00%",  0, _AU_REF["TRS"]),
        ("Conectividad Posventa",         "40,74%",  1, _AU_REF["Conectividad"]),
    ]},
    "0311Q VW": {"label":"Vara (0311Q VW) — VW Turismos", "total":6, "kpis":[
        ("Google Business Profile (GBP)", "100,00%", 2, _VW_REF["GBP"]),
        ("Medidas de Servicio (EC28)",    "0,97%",   2, _VW_REF["EC28"]),
        ("Conectividad Posventa",         "10,90%",  0, _VW_REF["Conectividad"]),
        ("Club Volkswagen Postventa",     "6,08%",   2, _VW_REF["Club VW"]),
    ]},
    "30070 VW": {"label":"Rivas VW (30070 VW) — VW Turismos", "total":6, "kpis":[
        ("Google Business Profile (GBP)", "100,00%", 2, _VW_REF["GBP"]),
        ("Medidas de Servicio (EC28)",    "0,50%",   2, _VW_REF["EC28"]),
        ("Conectividad Posventa",         "16,18%",  0, _VW_REF["Conectividad"]),
        ("Club Volkswagen Postventa",     "7,86%",   2, _VW_REF["Club VW"]),
    ]},
    "31523 VW": {"label":"Moncloa VW (31523 VW) — VW Turismos", "total":5, "kpis":[
        ("Google Business Profile (GBP)", "100,00%", 2, _VW_REF["GBP"]),
        ("Medidas de Servicio (EC28)",    "1,35%",   2, _VW_REF["EC28"]),
        ("Conectividad Posventa",         "7,32%",   0, _VW_REF["Conectividad"]),
        ("Club Volkswagen Postventa",     "4,58%",   1, _VW_REF["Club VW"]),
    ]},
    "30070 LCV": {"label":"Industriales (30070 LCV) — VW Comerciales", "total":3, "kpis":[
        ("Satisfacción General (CEM)",    "4,47",    0, _LCV_REF["Sat. Gral"]),
        ("Medidas de Servicio (EC28)",    "9,52%",   0, _LCV_REF["EC28"]),
        ("Conectividad Posventa",         "17,48%",  1, _LCV_REF["Conectividad"]),
        ("Servicios Rápidos (en el día)", "65,85%",  2, _LCV_REF["Servicios Rápidos"]),
    ]},
}

# Mapeo key → códigos CX de instalación
CX_CODES = {
    # Asesores
    "alvaro":             ("51AQ1 AU",),
    "estefania":          ("51AQ1 AU",),
    "codru":              ("51AQ1 AU",),
    "dennis":             ("51AQ1 AU",),
    "alberto_martinez":   ("51AQ2 AU",),
    "carlos_aguilar":     ("51AQ2 AU",),
    "magan":              ("51AQ3 AU",),
    "taboada":            ("51AQ3 AU",),
    "alejandro":          ("51AQ3 AU",),
    "jon":                ("51AQ3 AU",),
    "javier_diaz":        ("0311Q VW",),
    "jose_maria_vazquez": ("0311Q VW",),
    "juan":               ("0311Q VW",),
    "raul_munoz":         ("31523 AU", "31523 VW"),
    "jose_maria":         ("31523 AU", "31523 VW"),
    "guillermo":          ("30070 VW", "30070 LCV"),
    "javier_pina":        ("30070 VW", "30070 LCV"),
    "nuria":              ("30070 VW", "30070 LCV"),
    # Jefes mecánica
    "emilio_rivas_audi":       ("51AQ1 AU",),
    "alberto_martinez_ayala":  ("51AQ2 AU",),
    "izquierdo_canarias":      ("51AQ3 AU",),
    "jose_maria_moncloa_audi": ("31523 AU",),
    "jose_maria_moncloa_vw":   ("31523 VW",),
    "carlos_vara":             ("0311Q VW",),
    "fernando_rivas_vw":       ("30070 VW",),
    # Jefes chapa
    "luis_ramos_rivas_a_chapa": ("51AQ1 AU",),
    "luis_ramos_rivas_w_chapa": ("30070 VW",),
    "pericles_canarias_chapa":  ("51AQ3 AU",),
    "pericles_vara_chapa":      ("0311Q VW",),
    # Industriales
    "fernando_industriales":    ("30070 LCV",),
}

# Tareas pendientes de marzo (todas con grado vacío → se arrastran todas)
TAREAS_PREV = {
    "alberto_martinez":       ["Utilizar el Laser de desgaste","Full drive","Mirar condiciones de la copa horch. Esta en el portal interno","Kpis de one arriba indicados."],
    "alberto_martinez_ayala": ["Utilizar el Laser de desgaste.","Facturación de Recambios a dos años.","Deferencias y tablas.","Transporte alternativo","Productividad","Ocupacion.","Score One.","Service Cam"],
    "alejandro":              ["Utilizar el Laser de desgaste","Full drive","Mirar las condiciones de la copa horch","Conectividad"],
    "alvaro":                 ["Utilizar el Laser de desgaste","Mejorar en el CEM. Pedir el 5","Service Cam bajo","Puntos de calidad cem y connect","Health check"],
    "carlos_aguilar":         ["Utilizar el Laser de desgaste","Foco en desgaste","Mirar las condiciones de la copa horch. Están en el portal interno","Full drive","Service cam. Videos en one."],
    "carlos_vara":            ["Utilizar el Laser de desgaste","Conectividad","Puntos de calidad","Recogida y entrega","Productividad","Mantenimiento en 24 horas","Formacion"],
    "codru":                  ["Utilizar el Laser de desgaste","Full drive","Connect","Health check"],
    "dennis":                 ["Utilizar el Laser de desgaste","Service cam","Escalar puesto en copa horch","Sigue igual con la progresión que llevas en desgaste","Pedir 5 estrellas y conectar","Full long drive","Health check"],
    "emilio_rivas_audi":      ["Utilizar el Laser de desgaste","Puntos de Calidad","Indice de reclamaciones a clientes","Productividad","Ocupación Platon","Mantenimientos <24h","Score One.","Service Cam"],
    "estefania":              ["Utilizar el Laser de desgaste","Usar bien herramienta de ONE","Connect","Puntos de calidad","Mirar familias de desgaste","Health check"],
    "fernando_industriales":  ["Utilizar el Laser de desgaste","Radar y desgaste","Deferencias y demás","One Score","Servicio de Movilidad","Mantenimiento <24h","Diferidos"],
    "fernando_rivas_vw":      ["Utilizar el Laser de desgaste y radar","Ratio de desgaste","Variación de recambios","Variación de mano de obra","One Score"],
    "guillermo":              ["Utilizar el Laser de desgaste","Foco en discos","ONE score","Diferidos","Multimedia"],
    "izquierdo_canarias":     ["Utilizar el Laser de desgaste","Connect.","Bajo en deferencias, tablas","Mantenimiento <24 horas","Formación.","Score de ONE.","Service Cam ONE","Ojo mano de obra"],
    "javier_diaz":            ["Utilizar el Laser de desgaste","Club Volkswagen","Conectividad","Service cam","Diferidos"],
    "javier_pina":            ["Utilizar el Laser de desgaste","Foco en desgaste","Diferidos"],
    "jon":                    ["Utilizar el Laser de desgaste","Foco en desgaste","Mirar condiciones de la copa horch","Conectividad","Service cam","Diferidos"],
    "jose_maria":             ["Utilizar el Laser de desgaste"],
    "jose_maria_moncloa_audi":["Utilizar el Laser de desgaste","Desgaste general","Atenciones comerciales","One Score","Audi Move","Transporte alternativo","Ocupacióon","Mantenimiento <24 h","Formacion","Score ONE","Service Cam","Diferidos"],
    "jose_maria_moncloa_vw":  ["Utilizar el Laser de desgaste","Atenciones comerciales","Connect","Score de One","Puntos de Calidad","NPS","Movilidad","Ocupacion Platon","Mantenimiento <24 horas","Formacion","Service cam","Diferidos"],
    "jose_maria_vazquez":     ["Utilizar el Laser de desgaste","Mirar las condiciones la cum lauden","Conectividad","Club volkswagen"],
    "juan":                   ["Utilizar el Laser de desgaste","Conectividad","Club Volkswagen","Mirar las condiciones de la cum lauden"],
    "luis_ramos_rivas_a_chapa":["Mirar comentarios de Calidad en el portal","Chapa express","Reparacion de lunas","Registrar todos los kits de reparación","Recuperacion de siniestros"],
    "luis_ramos_rivas_w_chapa":["Venta cruzada de desgaste","Chapa express","Registrar todos los kits de reparación","Reparacion de lunas. Y también hacer cursos","Productividad","Tiempo de ciclo"],
    "magan":                  ["Utilizar el Laser de desgaste","Full drive","Mirar las condiciones de la copa horch","Conectividad","Service Cam"],
    "nuria":                  ["Utilizar el Laser de desgaste.","Sigue con la evolución que llevas en desgaste","Diferidos"],
    "pericles_canarias_chapa":["Chapa Express","Reparacion de Lunas","Meter todos los códigos de reparación","Recuperación de siniestros","Mínimo empatar en Junio en Mano de obra y recambios.","Productividad 120%","Tiempo de ciclo de carrocería.","Atentos a la formación"],
    "pericles_vara_chapa":    ["Chapa Express","Reparacion de lunas","Meter todos los códigos de reparación","Recuperacion de siniestros","Mínimo empatar en Junio en Mano de obra y recambios.","Productividad 120%","Tiempo de ciclo de carrocería.","Atentos a la formación"],
    "raul_munoz":             ["Utilizar el Laser de desgaste"],
    "taboada":                ["Utilizar el Laser de desgaste","Mejorar en todos los KPIS de One.","Foco en desgaste","Connect","Mirar condiciones de la copa horch","Service cam","Diferidos","Multimedia"],
}

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def cal1_pts(v):
    if v < 4:   return 0
    if v < 5:   return 10
    if v < 6:   return 20
    if v <= 7:  return 30
    return 50

def cal1_txt(v):
    pts = cal1_pts(v)
    val = f"{v:.2f}"
    base = f"CAL1SEM: {val} ({pts} puntos Iron Man)."
    if v >= 5:
        comt = {5:"Cumple el mínimo recomendado (5).",6:"Buen resultado.",7:"Resultado excelente."}.get(int(v),"Buen resultado.")
        return base + " " + comt, "positivo"
    return base + " Por debajo de 5. Revisar encuestas de satisfacción.", "mejorar"

def horch_txt(h, tipo="asesor"):
    if h is None:
        return "Horch Lauden: sin posición registrada en la competición de marca.", "neutro"
    if h <= 20:
        return f"Horch Lauden: posición {h}. Posición destacada (top 20).", "positivo"
    if h <= 50:
        return f"Horch Lauden: posición {h}. Posición aceptable.", "positivo"
    return f"Horch Lauden: posición {h}. Requiere atención.", "mejorar"

def one_bullets(d, media, label):
    pos, mej = [], []
    score_val = d["score"]
    str_score = str(score_val).replace(".",",")
    str_media_score = str(media["score"]).replace(".",",")
    if score_val > media["score"]:
        pos.append(f"Score SEM: {str_score} (media equipo: {str_media_score}) — por encima de la media")
    else:
        mej.append(f"Score SEM: {str_score} (media equipo: {str_media_score}) — por debajo de la media")

    cam_n = int(d["cam"].replace("%",""))
    if cam_n >= 8:
        pos.append(f"Service CAM Acum.: {d['cam']} (media equipo: {media['cam']}%; bueno ≥8%) — en rango positivo")
    elif cam_n >= 5:
        mej.append(f"Service CAM Acum.: {d['cam']} (media equipo: {media['cam']}%; bueno ≥8%) — en rango regular")
    else:
        mej.append(f"Service CAM Acum.: {d['cam']} (media equipo: {media['cam']}%; bueno ≥8%) — por debajo del objetivo")

    dif_n = int(d["dif"].replace("%",""))
    if dif_n >= 8:
        pos.append(f"Diferidos Acum.: {d['dif']} (media equipo: {media['dif']}%; bueno ≥8%) — en rango positivo")
    elif dif_n >= 5:
        mej.append(f"Diferidos Acum.: {d['dif']} (media equipo: {media['dif']}%; bueno ≥8%) — en rango regular")
    else:
        mej.append(f"Diferidos Acum.: {d['dif']} (media equipo: {media['dif']}%; bueno ≥8%) — por debajo del objetivo")

    multi_v = float(str(d["multi"]).replace(",","."))
    str_multi = str(d["multi"])
    str_media_multi = str(media["multi"]).replace(".",",")
    if multi_v >= 5:
        pos.append(f"Multimedia Acum.: {str_multi} (media equipo: {str_media_multi}; bueno ≥5) — en rango positivo")
    elif multi_v >= 3:
        mej.append(f"Multimedia Acum.: {str_multi} (media equipo: {str_media_multi}; bueno ≥5) — en rango regular")
    else:
        mej.append(f"Multimedia Acum.: {str_multi} (media equipo: {str_media_multi}; bueno ≥5) — por debajo del objetivo")

    return [f"Instalación ONE: {label}"] + pos, mej

def desgas_bullets_asesor(key):
    pos_rank, total, cats = DAS[key]
    pos, mej = [], []
    if total >= 200:
        pos.append(f"Desgaste total {total:.2f} pts (posición {pos_rank}/18 del ranking): por encima del objetivo de 200 pts.")
        cat_pos = " | ".join(f"{c}: {v:.2f} pts (media {MCAT[c]:.1f} pts) ✓" for c,v in cats.items() if v >= MCAT[c])
        cat_mej = " | ".join(f"{c}: {v:.2f} pts (media {MCAT[c]:.1f} pts)" for c,v in cats.items() if v < MCAT[c])
        if cat_pos:
            pos.append(f"Por encima de la media: {cat_pos}.")
        if cat_mej:
            mej.append(f"Por debajo de la media: {cat_mej}.")
    else:
        mej.append(f"Desgaste total {total:.2f} pts (posición {pos_rank}/18 del ranking): por debajo del objetivo de 200 pts. Seguimiento prioritario.")
        det = " | ".join(f"{c}: {v:.2f} pts" for c,v in cats.items())
        mej.append(f"Detalle por categoría: {det}.")
    return pos, mej

def fld_bullet(key):
    fld = IAS[key]["fld"]
    if fld > MFLD:
        return f"Full/Long Drive: {fld} operaciones (media equipo: {MFLD:.1f}). Por encima de la media.", "positivo"
    return f"Full/Long Drive: {fld} operaciones (media equipo: {MFLD:.1f}). Por debajo de la media.", "mejorar"

def fid_bullet(key):
    fid = IAS[key]["fid"]
    if fid is None:
        return None
    if fid >= MFID:
        return f"Fidelidad: {fid:.2f}% (media equipo: {MFID:.2f}%). Por encima o en línea con la media.", "positivo"
    return f"Fidelidad: {fid:.2f}% (media equipo: {MFID:.2f}%). Por debajo de la media.", "mejorar"

# ══════════════════════════════════════════════════════════════════════════════
# GENERADOR WORD
# ══════════════════════════════════════════════════════════════════════════════

def crear_word(titulo, subtitulo, secciones, path):
    doc = Document()
    s = doc.sections[0]
    s.page_width = Inches(8.27); s.page_height = Inches(11.69)
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1)

    h = doc.add_heading(titulo, level=0); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(subtitulo); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12); p.runs[0].font.color.rgb = RGBColor(0x44,0x44,0x44)
    f = doc.add_paragraph(f"Fecha: {HOY}"); f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f.runs[0].font.size = Pt(10); f.runs[0].font.color.rgb = RGBColor(0x88,0x88,0x88)
    doc.add_paragraph()

    for sec in secciones:
        doc.add_heading(sec["nombre"], level=1)

        # Texto introductorio (antes de positivos/mejoras)
        for txt in sec.get("intro", []):
            doc.add_paragraph(txt)

        if sec.get("positivo"):
            p2 = doc.add_paragraph()
            r = p2.add_run("Puntos positivos")
            r.bold = True; r.font.color.rgb = RGBColor(0x1A,0x7A,0x3C)
            for b in sec["positivo"]:
                doc.add_paragraph(b, style="List Bullet")

        if sec.get("mejorar"):
            p3 = doc.add_paragraph()
            r = p3.add_run("Puntos a mejorar")
            r.bold = True; r.font.color.rgb = RGBColor(0xC0,0x39,0x2B)
            for b in sec["mejorar"]:
                doc.add_paragraph(b, style="List Bullet")

        for blk in sec.get("cx_blocks", []):
            doc.add_paragraph(blk["label"])
            doc.add_paragraph(f"Total: {blk['total']}/8 puntos")
            if blk.get("positivo"):
                p2 = doc.add_paragraph(); r = p2.add_run("Puntos positivos")
                r.bold = True; r.font.color.rgb = RGBColor(0x1A, 0x7A, 0x3C)
                for b in blk["positivo"]:
                    doc.add_paragraph(b, style="List Bullet")
            if blk.get("mejorar"):
                p3 = doc.add_paragraph(); r = p3.add_run("Puntos a mejorar")
                r.bold = True; r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                for b in blk["mejorar"]:
                    doc.add_paragraph(b, style="List Bullet")

        if not sec.get("positivo") and not sec.get("mejorar") and "tareas" not in sec and not sec.get("cx_blocks"):
            doc.add_paragraph("Sin datos registrados")

        if "tareas" in sec:
            tbl = doc.add_table(rows=1, cols=5)
            tbl.style = "Table Grid"
            hdrs = tbl.rows[0].cells
            for i, h2 in enumerate(["Tarea","Responsable","Fecha límite","Observaciones","Grado de cumplimentación"]):
                hdrs[i].text = h2
                for r2 in hdrs[i].paragraphs[0].runs: r2.bold = True
            for t in sec["tareas"]:
                row = tbl.add_row().cells
                row[0].text = t; row[1].text = ""; row[2].text = ""; row[3].text = ""; row[4].text = ""

        doc.add_paragraph()

    doc.save(path)
    return path

# ══════════════════════════════════════════════════════════════════════════════
# CONSTRUCTORES DE SECCIONES
# ══════════════════════════════════════════════════════════════════════════════

SIN_DATOS = {"positivo":[], "mejorar":[]}

def sec_proactividad_asesor(key):
    pos, mej = desgas_bullets_asesor(key)
    txt, lado = fld_bullet(key)
    (pos if lado=="positivo" else mej).append(txt)
    res = fid_bullet(key)
    if res:
        txt2, lado2 = res
        (pos if lado2=="positivo" else mej).append(txt2)
    return {"nombre":"Proactividad comercial","positivo":pos,"mejorar":mej}

def sec_calidad_asesor(key):
    txt, lado = cal1_txt(IAS[key]["cal1"])
    pos = [txt] if lado == "positivo" else []
    mej = [txt] if lado == "mejorar" else []
    return {"nombre":"Calidad","positivo":pos,"mejorar":mej}

def sec_horch_asesor(key):
    txt, _ = horch_txt(IAS[key]["horch"])
    h = IAS[key]["horch"]
    pos = [txt] if (h is None or h <= 50) else []
    mej = [txt] if (h is not None and h > 50) else []
    return {"nombre":"Horch Lauden","positivo":pos,"mejorar":mej}

def sec_one_asesor(key):
    if key not in OAS:
        return {"nombre":"ONE KVPS","positivo":[],"mejorar":[]}
    pos, mej = one_bullets(OAS[key], ONE_AM, "Asesor")
    return {"nombre":"ONE KVPS","positivo":pos,"mejorar":mej}

def _cx_block(code):
    d = CX_APR[code]
    pos, mej = [], []
    for name, val, pts, ref in d["kpis"]:
        bullet = f"{name}: {val} — {pts}/2 pts  |  Referencia: {ref}"
        (pos if pts == 2 else mej).append(bullet)
    return {"label": d["label"], "total": d["total"], "positivo": pos, "mejorar": mej}

def sec_cx_for(key):
    codes = CX_CODES.get(key, ())
    if not codes:
        return {"nombre": "Cuadro de Calidad CX 2026", "positivo": [], "mejorar": []}
    return {"nombre": "Cuadro de Calidad CX 2026", "cx_blocks": [_cx_block(c) for c in codes]}

def sec_one_inst(inst_key, label_extra=""):
    if inst_key not in OINST:
        return {"nombre":"ONE KVPS","positivo":[],"mejorar":[]}
    d = OINST[inst_key]
    pos, mej = one_bullets(d, ONE_IM, d["label"])
    return {"nombre":"ONE KVPS","positivo":pos,"mejorar":mej}

def sec_proactividad_jt_mecanica(key):
    d = IJT[key]
    pos, mej = [], []
    # Desgaste
    if key in DIT:
        pos_rank, total, cats = DIT[key]
        if total >= 220:
            pos.append(f"Desgaste total {total:.2f}% (posición {pos_rank}/7 del ranking): por encima del umbral de referencia (220%).")
        else:
            mej.append(f"Desgaste total {total:.2f}% (posición {pos_rank}/7 del ranking): por debajo del umbral de referencia (220%). Seguimiento prioritario.")
        cat_pos = " | ".join(f"{c}: {v:.2f}% (media {MCAT_INST[c]:.2f}%) ✓" for c, v in cats.items() if v >= MCAT_INST[c])
        cat_mej = " | ".join(f"{c}: {v:.2f}% (media {MCAT_INST[c]:.2f}%)" for c, v in cats.items() if v < MCAT_INST[c])
        if cat_pos:
            pos.append(f"Por encima de la media: {cat_pos}.")
        if cat_mej:
            mej.append(f"Por debajo de la media: {cat_mej}.")
    else:
        mej.append("Desgaste: pendiente de ranking Desgaste Instalaciones 2026-04.")
    mej_note = d.get("desgas_note")
    if mej_note:
        mej.append(mej_note)
    # Variaciones
    vr = d["var_rec"]; vm = d["var_mo"]
    (pos if vr >= 0 else mej).append(f"Variación recambios respecto a hace 2 años: {'+' if vr>=0 else ''}{vr:.2f}%.")
    (pos if vm >= 0 else mej).append(f"Variación mano de obra respecto a hace 2 años: {'+' if vm>=0 else ''}{vm:.2f}%.")
    mej.append("ICC Proactividad: pendiente de ranking ICC Instalaciones 2026-04.")
    return {"nombre":"Proactividad comercial","positivo":pos,"mejorar":mej}

def sec_proactividad_jt_chapa(key):
    d = IJT[key]
    pos, mej = [], []
    vr = d["var_rec"]; vm = d["var_mo"]
    (pos if vr >= 0 else mej).append(f"Variación recambios respecto a hace 2 años: {'+' if vr>=0 else ''}{vr:.2f}%.")
    (pos if vm >= 0 else mej).append(f"Variación mano de obra respecto a hace 2 años: {'+' if vm>=0 else ''}{vm:.2f}%.")
    mej.append("ICC Proactividad: pendiente de ranking ICC Instalaciones 2026-04.")
    return {"nombre":"Proactividad comercial","positivo":pos,"mejorar":mej}

def sec_calidad_jt(key):
    txt, lado = cal1_txt(IJT[key]["cal1"])
    pos = [txt] if lado == "positivo" else []
    mej = [txt] if lado == "mejorar" else []
    mej.append("ICC Calidad: pendiente de ranking ICC Instalaciones 2026-04.")
    return {"nombre":"Calidad","positivo":pos,"mejorar":mej}

def sec_procesos_jt(key):
    d = IJT[key]
    prod = d.get("prod"); obj = d.get("obj_prod",93)
    if prod is None:
        return {"nombre":"Procesos","positivo":[],"mejorar":[]}
    gap = prod - obj
    txt = f"Productividad: {prod:.2f}% sobre objetivo {obj}% ({'+' if gap>=0 else ''}{gap:.2f}pp). {'Objetivo cumplido.' if gap>=0 else 'Objetivo no alcanzado.'}"
    pos = [txt] if gap >= 0 else []
    mej = [txt] if gap < 0 else []
    return {"nombre":"Procesos","positivo":pos,"mejorar":mej}

def sec_rankings_jt(key):
    d = IJT[key]
    pos, mej = [], []
    icc = d["icc_pos"]
    pos.append(f"ICC global: posición {icc} | Iron Man ICC: posición {icc} en el ranking de marca. {'Posición destacada (top 10).' if icc<=10 else 'Posición aceptable.' if icc<=30 else 'Posición a mejorar.'}")
    txt_h, lado_h = horch_txt(d["horch"], "jt")
    (pos if lado_h in ("positivo","neutro") else mej).append(txt_h)
    return {"nombre":"Rankings de marca","positivo":pos,"mejorar":mej}

# ══════════════════════════════════════════════════════════════════════════════
# DEFINICIÓN DE INFORMES
# ══════════════════════════════════════════════════════════════════════════════

def secs_sin_datos(*nombres):
    return [{"nombre":n,"positivo":[],"mejorar":[]} for n in nombres]

def tareas_sec(key):
    return {"nombre":"Tareas pendientes","tareas": TAREAS_PREV.get(key,[])}

def informe_asesor(key, titulo, subtitulo):
    secs = [
        sec_proactividad_asesor(key),
        {"nombre":"Digital","positivo":[],"mejorar":[]},
        sec_calidad_asesor(key),
        {"nombre":"Procesos","positivo":[],"mejorar":[]},
        {"nombre":"Recursos generales","positivo":[],"mejorar":[]},
        sec_horch_asesor(key),
        sec_cx_for(key),
        sec_one_asesor(key),
        tareas_sec(key),
    ]
    fname = f"informe_seguimiento_{key}_{PERIODO}.docx"
    path = os.path.join(OUTPUT, fname)
    crear_word(titulo, subtitulo, secs, path)
    return fname

def informe_jt_mecanica(key, titulo, subtitulo, one_key):
    secs = [
        sec_proactividad_jt_mecanica(key),
        {"nombre":"Digital","positivo":[],"mejorar":[]},
        sec_calidad_jt(key),
        sec_procesos_jt(key),
        {"nombre":"Recursos generales","positivo":[],"mejorar":[]},
        {"nombre":"Carrocería","positivo":[],"mejorar":[]},
        {"nombre":"Personas","positivo":[],"mejorar":[]},
        sec_rankings_jt(key),
        sec_cx_for(key),
        sec_one_inst(one_key),
        tareas_sec(key),
    ]
    fname = f"informe_seguimiento_{key}_{PERIODO}.docx"
    crear_word(titulo, subtitulo, secs, os.path.join(OUTPUT, fname))
    return fname

def informe_jt_chapa(key, titulo, subtitulo):
    secs = [
        sec_proactividad_jt_chapa(key),
        {"nombre":"Carrocería","positivo":[],"mejorar":[]},
        {"nombre":"Personas","positivo":[],"mejorar":[]},
        sec_rankings_jt(key),
        sec_cx_for(key),
        tareas_sec(key),
    ]
    fname = f"informe_seguimiento_{key}_{PERIODO}.docx"
    crear_word(titulo, subtitulo, secs, os.path.join(OUTPUT, fname))
    return fname

def informe_jt_industriales(key, titulo, subtitulo):
    # Solo ICC (no disponible) + ONE + Tareas
    secs = [
        {"nombre":"Proactividad comercial","positivo":[],"mejorar":[]},
        {"nombre":"Digital","positivo":[],"mejorar":[]},
        {"nombre":"Calidad","positivo":[],"mejorar":[]},
        {"nombre":"Recursos generales","positivo":[],"mejorar":[]},
        {"nombre":"Carrocería","positivo":[],"mejorar":[]},
        {"nombre":"Personas","positivo":[],"mejorar":[]},
        sec_cx_for(key),
        sec_one_inst("rivas_vw"),  # 30070 mismo código
        tareas_sec(key),
    ]
    fname = f"informe_seguimiento_{key}_{PERIODO}.docx"
    crear_word(titulo, subtitulo, secs, os.path.join(OUTPUT, fname))
    return fname

# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    generados = []

    # ── ASESORES ──────────────────────────────────────────────────────────────
    asesores = [
        ("alvaro",            "Seguimiento Rivas Audi — Abril 2026", "Álvaro · Rivas Audi"),
        ("estefania",         "Seguimiento Rivas Audi — Abril 2026", "Estefanía · Rivas Audi"),
        ("codru",             "Seguimiento Rivas Audi — Abril 2026", "Alex Codru · Rivas Audi"),
        ("dennis",            "Seguimiento Rivas Audi — Abril 2026", "Dennis · Rivas Audi"),
        ("alberto_martinez",  "Seguimiento Ayala — Abril 2026",      "Alberto Martínez · Ayala"),
        ("carlos_aguilar",    "Seguimiento Ayala — Abril 2026",      "Carlos Aguilar · Ayala"),
        ("magan",             "Seguimiento Canarias — Abril 2026",   "Magán · Canarias"),
        ("taboada",           "Seguimiento Canarias — Abril 2026",   "Taboada · Canarias"),
        ("alejandro",         "Seguimiento Canarias — Abril 2026",   "Alejandro · Canarias"),
        ("jon",               "Seguimiento Canarias — Abril 2026",   "Jon · Canarias"),
        ("javier_diaz",       "Seguimiento Vara — Abril 2026",       "Javier Díaz Mesa · Vara"),
        ("jose_maria_vazquez","Seguimiento Vara — Abril 2026",       "José María Vázquez · Vara"),
        ("juan",              "Seguimiento Vara — Abril 2026",       "Juan · Vara"),
        ("raul_munoz",        "Seguimiento Moncloa — Abril 2026",    "Raúl Muñoz · Moncloa"),
        ("jose_maria",        "Seguimiento Moncloa — Abril 2026",    "José María Campos · Moncloa"),
        ("guillermo",         "Seguimiento Rivas VW — Abril 2026",   "Guillermo · Rivas VW / Industriales"),
        ("javier_pina",       "Seguimiento Rivas VW — Abril 2026",   "Javier Pina · Rivas VW / Industriales"),
        ("nuria",             "Seguimiento Rivas VW — Abril 2026",   "Nuria · Rivas VW / Industriales"),
    ]
    for key, titulo, sub in asesores:
        f = informe_asesor(key, titulo, sub)
        generados.append(f)
        print(f"  OK  {f}")

    # ── JEFES MECÁNICA ────────────────────────────────────────────────────────
    jt_mec = [
        ("emilio_rivas_audi",       "Seguimiento Rivas Audi — Abril 2026",  "Emilio · Rivas Audi",          "rivas_audi"),
        ("alberto_martinez_ayala",  "Seguimiento Ayala — Abril 2026",       "Alberto Martínez · Ayala",     "ayala"),
        ("carlos_vara",             "Seguimiento Vara — Abril 2026",        "Carlos · Vara",                "vara"),
        ("izquierdo_canarias",      "Seguimiento Canarias — Abril 2026",    "Izquierdo · Canarias",         "canarias"),
        ("jose_maria_moncloa_audi", "Seguimiento Moncloa Audi — Abril 2026","José María Campos · Moncloa Audi","moncloa"),
        ("jose_maria_moncloa_vw",   "Seguimiento Moncloa VW — Abril 2026",  "José María Campos · Moncloa VW",  "moncloa"),
        ("fernando_rivas_vw",       "Seguimiento Rivas VW — Abril 2026",    "Fernando · Rivas VW",          "rivas_vw"),
    ]
    for key, titulo, sub, one_k in jt_mec:
        f = informe_jt_mecanica(key, titulo, sub, one_k)
        generados.append(f)
        print(f"  OK  {f}")

    # ── JEFES CHAPA ───────────────────────────────────────────────────────────
    jt_chapa = [
        ("luis_ramos_rivas_a_chapa", "Seguimiento Rivas Audi Chapa — Abril 2026", "Luis Ramos · Rivas Audi Chapa"),
        ("luis_ramos_rivas_w_chapa", "Seguimiento Rivas VW Chapa — Abril 2026",   "Luis Ramos · Rivas VW Chapa"),
        ("pericles_canarias_chapa",  "Seguimiento Canarias Chapa — Abril 2026",   "Pericles · Canarias Chapa"),
        ("pericles_vara_chapa",      "Seguimiento Vara Chapa — Abril 2026",       "Pericles · Vara Chapa"),
    ]
    for key, titulo, sub in jt_chapa:
        f = informe_jt_chapa(key, titulo, sub)
        generados.append(f)
        print(f"  OK  {f}")

    # ── INDUSTRIALES ──────────────────────────────────────────────────────────
    f = informe_jt_industriales("fernando_industriales",
                                 "Seguimiento Industriales — Abril 2026",
                                 "Fernando · Industriales")
    generados.append(f)
    print(f"  OK  {f}")

    print(f"\nGenerados: {len(generados)}/30")

if __name__ == "__main__":
    main()
