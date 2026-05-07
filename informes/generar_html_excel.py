#!/usr/bin/env /opt/homebrew/bin/python3
"""
Genera HTML y Excel multi-pestaña para cada persona de seguimiento.
Por cada persona: una pestaña/hoja por mes + Evolución + Objetivos.
Lee todos los Word de esa persona y los agrupa.
"""

import glob
import os
import re
import datetime
from docx import Document

from openpyxl import Workbook
from openpyxl.styles import (
    PatternFill, Font, Alignment, Border, Side
)
from openpyxl.utils import get_column_letter

# ══════════════════════════════════════════════════════════════════════════════
# CONSTANTES
# ══════════════════════════════════════════════════════════════════════════════

MESES_ES = {
    "01":"Enero","02":"Febrero","03":"Marzo","04":"Abril",
    "05":"Mayo","06":"Junio","07":"Julio","08":"Agosto",
    "09":"Septiembre","10":"Octubre","11":"Noviembre","12":"Diciembre"
}

JT_MECANICA = {
    "emilio_rivas_audi","alberto_martinez_ayala","carlos_vara",
    "izquierdo_canarias","jose_maria_moncloa_audi","jose_maria_moncloa_vw",
    "fernando_rivas_vw"
}

OBJ_PROD_ESPECIFICO = {
    "alberto_martinez_ayala": 105,
    "emilio_rivas_audi": 93,
    "carlos_vara": 93,
    "izquierdo_canarias": 93,
    "jose_maria_moncloa_audi": 93,
    "jose_maria_moncloa_vw": 93,
    "fernando_rivas_vw": 89,
    "luis_ramos_rivas_a_chapa": 120,
    "luis_ramos_rivas_w_chapa": 120,
    "pericles_canarias_chapa": 120,
    "pericles_vara_chapa": 120,
}

# (nombre_métrica, regex, higher_is_better)
EVO_METRICS_ASESOR = [
    ("Desgaste total",   r'Desgaste total ([\d,.]+)\s*(?:pts|%)',            True),
    ("Posición ranking", r'posición (\d+)/\d+',                              False),
    ("CAL1SEM",          r'CAL1SEM:\s*([\d,.]+)',                            True),
    ("Full/Long Drive",  r'Full/Long Drive:\s*(\d+)\s*operaciones',          True),
    ("Fidelidad (%)",    r'Fidelidad:\s*([\d,.]+)%',                        True),
    ("Horch Lauden",     r'Horch Lauden:\s*posici[oó]n\s*(\d+)',             False),
    ("CX total (/8)",   r'Total:\s*(\d+)/8\s*puntos',                      True),
    ("Score ONE",        r'Score SEM:\s*([\d,]+)',                           True),
    ("Service CAM (%)",  r'Service CAM Acum\.:\s*(\d+)%',                   True),
    ("Diferidos (%)",    r'Diferidos Acum\.:\s*(\d+)%',                     True),
    ("Multimedia",       r'Multimedia Acum\.:\s*([\d,]+)',                   True),
]

_ICC_MOD_PUNT = r'ICC M[oó]dulo {mod}.+?Puntuaci[oó]n del m[oó]dulo:\s*([\d,]+)'

def _icc(mod):
    return _ICC_MOD_PUNT.replace("{mod}", mod)

EVO_METRICS_JT_MECANICA = [
    ("Desgaste total (%)", r'Desgaste total ([\d,.]+)\s*%',                  True),
    ("Posición ranking",   r'posici[oó]n (\d+)/\d+',                          False),
    ("CAL1SEM",            r'CAL1SEM:\s*([\d,.]+)',                          True),
    ("Productividad (%)",  r'Productividad:\s*([\d,.]+)%',                   True),
    ("Var Rec 2A (%)",     r'Variaci[oó]n recambios.+?([+\-]?[\d,.]+)%',    True),
    ("Var MO 2A (%)",      r'Variaci[oó]n mano de obra.+?([+\-]?[\d,.]+)%', True),
    ("Horch Lauden",       r'Horch Lauden:\s*posici[oó]n\s*(\d+)',           False),
    ("CX total (/8)",     r'Total:\s*(\d+)/8\s*puntos',                     True),
    # ICC módulos
    ("ICC posición global",  r'ICC global:\s*posici[oó]n\s*(\d+)',           False),
    ("ICC Proactividad",     _icc("Proactividad"),                           True),
    ("ICC Digital",          _icc("Digital"),                                True),
    ("ICC Calidad",          _icc("Calidad"),                                True),
    ("ICC Recursos",         _icc("Recursos"),                               True),
    ("ICC Carrocería",       _icc("Carrocer"),                               True),
    ("ICC Personas",         _icc("Personas"),                               True),
    ("CX total (/8)",       r'Total:\s*(\d+)/8\s*puntos',                   True),
    # ONE
    ("Score ONE",          r'Score SEM:\s*([\d,]+)',                         True),
    ("Service CAM (%)",    r'Service CAM Acum\.:\s*(\d+)%',                 True),
    ("Diferidos (%)",      r'Diferidos Acum\.:\s*(\d+)%',                   True),
    ("Multimedia",         r'Multimedia Acum\.:\s*([\d,]+)',                 True),
]

EVO_METRICS_JT_CHAPA = [
    ("CAL1SEM",            r'CAL1SEM:\s*([\d,.]+)',                          True),
    ("Productividad (%)",  r'Productividad:\s*([\d,.]+)%',                   True),
    ("Var Rec 2A (%)",     r'Variaci[oó]n recambios.+?([+\-]?[\d,.]+)%',    True),
    ("Var MO 2A (%)",      r'Variaci[oó]n mano de obra.+?([+\-]?[\d,.]+)%', True),
    ("Horch Lauden",       r'Horch Lauden:\s*posici[oó]n\s*(\d+)',           False),
    ("CX total (/8)",     r'Total:\s*(\d+)/8\s*puntos',                     True),
    # ICC módulos aplicables a chapa
    ("ICC posición global",  r'ICC global:\s*posici[oó]n\s*(\d+)',           False),
    ("ICC Carrocería",       _icc("Carrocer"),                               True),
    ("ICC Personas",         _icc("Personas"),                               True),
]

EVO_METRICS_INDUSTRIALES = [
    # ICC módulos (Fernando Industriales solo tiene datos ICC)
    ("ICC posición global",  r'ICC global:\s*posici[oó]n\s*(\d+)',           False),
    ("ICC Proactividad",     _icc("Proactividad"),                           True),
    ("ICC Digital",          _icc("Digital"),                                True),
    ("ICC Calidad",          _icc("Calidad"),                                True),
    ("ICC Recursos",         _icc("Recursos"),                               True),
    ("ICC Carrocería",       _icc("Carrocer"),                               True),
    ("ICC Personas",         _icc("Personas"),                               True),
    ("CX total (/8)",       r'Total:\s*(\d+)/8\s*puntos',                   True),
    # ONE
    ("Score ONE",          r'Score SEM:\s*([\d,]+)',                         True),
    ("Service CAM (%)",    r'Service CAM Acum\.:\s*(\d+)%',                 True),
    ("Diferidos (%)",      r'Diferidos Acum\.:\s*(\d+)%',                   True),
    ("Multimedia",         r'Multimedia Acum\.:\s*([\d,]+)',                 True),
]

# Objetivo por métrica y tipo — aparece como última fila en la tabla de evolución
EVO_OBJ = {
    "asesor": {
        "Desgaste total":   "220 pts",
        "Posición ranking": "—",
        "CAL1SEM":          "≥6",
        "Full/Long Drive":  "11,1 ops",
        "Fidelidad (%)":    "63,48%",
        "Horch Lauden":     "<20",
        "CX total (/8)":   "≥7",
        "Score ONE":        "6,2",
        "Service CAM (%)":  "8%",
        "Diferidos (%)":    "15%",
        "Multimedia":       "5",
    },
    "jt_mecanica": {
        "Desgaste total (%)":  ">220%",
        "Posición ranking":    "—",
        "CAL1SEM":             "≥6",
        "Productividad (%)":   "93%",   # sobrescrito por instalación si aplica
        "Var Rec 2A (%)":      ">8%",
        "Var MO 2A (%)":       ">6%",
        "Horch Lauden":        "<20",
        "CX total (/8)":      "≥7",
        "ICC posición global": "—",
        "ICC Proactividad":    "—",
        "ICC Digital":         "—",
        "ICC Calidad":         "—",
        "ICC Recursos":        "—",
        "ICC Carrocería":      "—",
        "ICC Personas":        "—",
        "Score ONE":           "7,3",
        "Service CAM (%)":     "8%",
        "Diferidos (%)":       "15%",
        "Multimedia":          "5",
    },
    "jt_chapa": {
        "CAL1SEM":             "≥6",
        "Productividad (%)":   "120%",
        "Var Rec 2A (%)":      ">8%",
        "Var MO 2A (%)":       ">6%",
        "Horch Lauden":        "<20",
        "CX total (/8)":      "≥7",
        "ICC posición global": "—",
        "ICC Carrocería":      "—",
        "ICC Personas":        "—",
    },
    "industriales": {
        "ICC posición global": "—",
        "ICC Proactividad":    "—",
        "ICC Digital":         "—",
        "ICC Calidad":         "—",
        "ICC Recursos":        "—",
        "ICC Carrocería":      "—",
        "ICC Personas":        "—",
        "CX total (/8)":      "≥7",
        "Score ONE":           "7,3",
        "Service CAM (%)":     "8%",
        "Diferidos (%)":       "15%",
        "Multimedia":          "5",
    },
}

# Bloques de evolución — agrupan las métricas por sección del informe
EVO_BLOCKS = {
    "asesor": [
        ("Proactividad comercial", ["Desgaste total", "Posición ranking", "Full/Long Drive", "Fidelidad (%)"]),
        ("Calidad",                ["CAL1SEM"]),
        ("Horch Lauden",           ["Horch Lauden"]),
        ("Cuadro CX",              ["CX total (/8)"]),
        ("ONE KVPS",               ["Score ONE", "Service CAM (%)", "Diferidos (%)", "Multimedia"]),
    ],
    "jt_mecanica": [
        ("Proactividad comercial", ["Desgaste total (%)", "Posición ranking", "Var Rec 2A (%)", "Var MO 2A (%)", "ICC Proactividad"]),
        ("Digital",                ["ICC Digital"]),
        ("Calidad",                ["CAL1SEM", "ICC Calidad"]),
        ("Procesos",               ["Productividad (%)"]),
        ("Recursos generales",     ["ICC Recursos"]),
        ("Carrocería",             ["ICC Carrocería"]),
        ("Personas",               ["ICC Personas"]),
        ("Rankings de marca",      ["ICC posición global", "Horch Lauden"]),
        ("Cuadro CX",              ["CX total (/8)"]),
        ("ONE KVPS",               ["Score ONE", "Service CAM (%)", "Diferidos (%)", "Multimedia"]),
    ],
    "jt_chapa": [
        ("Proactividad comercial", ["Var Rec 2A (%)", "Var MO 2A (%)"]),
        ("Calidad",                ["CAL1SEM"]),
        ("Procesos",               ["Productividad (%)"]),
        ("Carrocería",             ["ICC Carrocería"]),
        ("Personas",               ["ICC Personas"]),
        ("Rankings de marca",      ["ICC posición global", "Horch Lauden"]),
        ("Cuadro CX",              ["CX total (/8)"]),
    ],
    "industriales": [
        ("Proactividad comercial", ["ICC Proactividad"]),
        ("Digital",                ["ICC Digital"]),
        ("Calidad",                ["ICC Calidad"]),
        ("Recursos generales",     ["ICC Recursos"]),
        ("Carrocería",             ["ICC Carrocería"]),
        ("Personas",               ["ICC Personas"]),
        ("Rankings de marca",      ["ICC posición global"]),
        ("Cuadro CX",              ["CX total (/8)"]),
        ("ONE KVPS",               ["Score ONE", "Service CAM (%)", "Diferidos (%)", "Multimedia"]),
    ],
}

OBJETIVOS = {
    "asesor": [
        ("Desgaste total",     ">200 pts",              "≥200 pts",    "<200 pts — seguimiento prioritario"),
        ("CAL1SEM",            "≥5,0",                  "≥5,0",        "<5,0"),
        ("Full/Long Drive",    "Media equipo: 11,1 ops",">11,1 ops",   "≤11,1 ops"),
        ("Fidelidad",          "Media equipo: 63,48%",  "≥63,48%",     "<63,48%"),
        ("Horch Lauden (pos)", "≤50 aceptable",         "≤20 dest. / ≤50 acept.", ">50 — requiere atención"),
        ("ONE Score SEM",      "Por encima de la media (6,2)", ">6,2", "≤6,2"),
        ("ONE Service CAM",    "≥8%",                   "≥8%",         "5-8% regular / <5% mejorar"),
        ("ONE Diferidos",      "≥8%",                   "≥8%",         "5-8% regular / <5% mejorar"),
        ("ONE Multimedia",     "≥5",                    "≥5",          "3-5 regular / <3 mejorar"),
    ],
    "jt_mecanica": [
        ("Desgaste total (%)", ">220%",                 "≥220%",       "<220% — seguimiento prioritario"),
        ("CAL1SEM",            "≥5,0",                  "≥5,0",        "<5,0"),
        ("Productividad",      "Específico por instalación", "≥objetivo", "<objetivo"),
        ("Var Rec 2A",         "≥0%",                   "≥0%",         "<0%"),
        ("Var MO 2A",          "≥0%",                   "≥0%",         "<0%"),
        ("Horch Lauden (pos)", "≤50 aceptable",         "≤20 dest. / ≤50 acept.", ">50 — requiere atención"),
        ("ONE Score SEM",      "Por encima de la media (7,3)", ">7,3", "≤7,3"),
        ("ONE Service CAM",    "≥8%",                   "≥8%",         "5-8% regular / <5% mejorar"),
        ("ONE Diferidos",      "≥8%",                   "≥8%",         "5-8% regular / <5% mejorar"),
        ("ONE Multimedia",     "≥5",                    "≥5",          "3-5 regular / <3 mejorar"),
    ],
    "jt_chapa": [
        ("CAL1SEM",            "≥5,0",                  "≥5,0",        "<5,0"),
        ("Productividad",      "120%",                  "≥120%",       "<120%"),
        ("Var Rec 2A",         "≥0%",                   "≥0%",         "<0%"),
        ("Var MO 2A",          "≥0%",                   "≥0%",         "<0%"),
        ("Horch Lauden (pos)", "≤50 aceptable",         "≤20 dest. / ≤50 acept.", ">50 — requiere atención"),
    ],
    "industriales": [
        ("ONE Score SEM",      "Por encima de la media (7,3)", ">7,3", "≤7,3"),
        ("ONE Service CAM",    "≥8%",                   "≥8%",         "5-8% regular / <5% mejorar"),
        ("ONE Diferidos",      "≥8%",                   "≥8%",         "5-8% regular / <5% mejorar"),
        ("ONE Multimedia",     "≥5",                    "≥5",          "3-5 regular / <3 mejorar"),
    ],
}

# ══════════════════════════════════════════════════════════════════════════════
# PARSER DEL WORD (sin cambios)
# ══════════════════════════════════════════════════════════════════════════════

ICC_MOD_RE = re.compile(
    r"ICC\s+M[oó]dulo\s+(.+?)\s*[—\-]\s*Posici[oó]n\s+(\d+)\s+de\s+(\d+)"
    r"(?:\s*\|\s*Puntuaci[oó]n del m[oó]dulo:\s*([\d,\.]+))?",
    re.I
)

def parse_word(path):
    doc = Document(path)
    meta = {"titulo": "", "subtitulo": "", "fecha": ""}
    sections = []
    cur_sec = None
    bloque = None

    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if not txt:
            continue
        sname = p.style.name
        if sname == "Title" or (i == 0 and not meta["titulo"]):
            meta["titulo"] = txt
        elif sname == "Normal" and i == 1 and not meta["subtitulo"]:
            meta["subtitulo"] = txt
        elif sname == "Normal" and txt.startswith("Fecha:") and not meta["fecha"]:
            meta["fecha"] = txt
        elif sname == "Heading 1":
            cur_sec = {"nombre": txt, "positivo": [], "mejorar": [], "neutro": []}
            sections.append(cur_sec)
            bloque = None
        elif txt == "Puntos positivos":
            bloque = "positivo"
        elif txt == "Puntos a mejorar":
            bloque = "mejorar"
        elif txt == "Sin datos registrados":
            if cur_sec is not None:
                cur_sec["sin_datos"] = True
        elif cur_sec is not None and bloque and sname in ("List Bullet", "List Paragraph", "Normal"):
            cur_sec[bloque].append(txt)
        elif cur_sec is not None and not bloque and sname == "Normal":
            cur_sec["neutro"].append(txt)

    tabla_tareas = None
    for tbl in doc.tables:
        filas = [[c.text.strip() for c in row.cells] for row in tbl.rows]
        if filas and "Tarea" in filas[0]:
            tabla_tareas = filas
            break

    return meta, sections, tabla_tareas


def parse_icc_bullet(txt):
    if not re.match(r'^ICC\s*-\s*', txt, re.I):
        return None
    rest = re.sub(r'^ICC\s*-\s*', '', txt, flags=re.I).strip()
    colon = rest.find(':')
    if colon == -1:
        return None
    indicador = rest[:colon].strip()
    after = rest[colon + 1:].strip()
    m_media = re.search(r'\(media\s+', after, re.I)
    if m_media:
        valor = after[:m_media.start()].strip()
        inner_start = m_media.end()
        close = after.find(')', inner_start)
        if close == -1:
            close = len(after)
        inner = after[inner_start:close]
        m_top = re.search(r',\s*Top20\s+', inner, re.I)
        if m_top:
            media = inner[:m_top.start()].strip()
            top20 = inner[m_top.end():].strip()
        else:
            media = inner.strip()
            top20 = ""
        comment = after[close + 1:].strip().lstrip('.')
    else:
        valor = after
        media = ""
        top20 = ""
        comment = ""
    return indicador, valor, media, top20, comment


def parse_icc_module_header(txt):
    m = ICC_MOD_RE.match(txt)
    if not m:
        return None
    return m.group(1).strip(), m.group(2), m.group(3), (m.group(4) or "").strip()


# ══════════════════════════════════════════════════════════════════════════════
# AGRUPACIÓN Y DETECCIÓN DE TIPO
# ══════════════════════════════════════════════════════════════════════════════

RE_FNAME = re.compile(r'informe_seguimiento_(.+?)_(\d{4}-\d{2})\.docx$')

def group_files(base):
    """Devuelve {key: [(periodo, path), ...]} ordenado por periodo."""
    pattern = os.path.join(base, "informe_seguimiento_*.docx")
    groups = {}
    for f in sorted(glob.glob(pattern)):
        bn = os.path.basename(f)
        if bn.startswith("~") or "_CEX_" in bn:
            continue
        m = RE_FNAME.search(bn)
        if not m:
            continue
        key, periodo = m.group(1), m.group(2)
        groups.setdefault(key, []).append((periodo, f))
    return groups


def detect_tipo(key):
    if "industriales" in key:
        return "industriales"
    if "_chapa" in key:
        return "jt_chapa"
    if key in JT_MECANICA:
        return "jt_mecanica"
    return "asesor"


def periodo_label(periodo):
    """'2026-04' → 'Abr 2026'"""
    parts = periodo.split("-")
    if len(parts) == 2:
        mes = MESES_ES.get(parts[1], parts[1])[:3]
        return f"{mes} {parts[0]}"
    return periodo


def evo_metrics_for_tipo(tipo):
    if tipo == "asesor":
        return EVO_METRICS_ASESOR
    if tipo == "jt_mecanica":
        return EVO_METRICS_JT_MECANICA
    if tipo == "jt_chapa":
        return EVO_METRICS_JT_CHAPA
    return EVO_METRICS_INDUSTRIALES


def extract_evolution_metrics(sections, tipo):
    """Extrae métricas clave de los bullets ya parseados del Word."""
    all_text = []
    for sec in sections:
        for b in sec.get("neutro", []):
            all_text.append(b)
        for b in sec.get("positivo", []):
            all_text.append(b)
        for b in sec.get("mejorar", []):
            all_text.append(b)
    combined = " | ".join(all_text)

    metrics = {}
    for name, pattern, higher_better in evo_metrics_for_tipo(tipo):
        m = re.search(pattern, combined, re.I)
        if m:
            raw = m.group(1).replace(",", ".").replace("+", "")
            try:
                metrics[name] = float(raw)
            except ValueError:
                metrics[name] = raw
    return metrics


# ══════════════════════════════════════════════════════════════════════════════
# HTML — HELPERS
# ══════════════════════════════════════════════════════════════════════════════

CSS = """
* { box-sizing: border-box; margin: 0; padding: 0; }
body { font-family: 'Segoe UI', Arial, sans-serif; background: #F0F2F5; color: #222; }
.page { max-width: 960px; margin: 0 auto; background: #fff; box-shadow: 0 2px 12px rgba(0,0,0,.15); }

/* Header */
.header { background: #1A2B4A; color: #fff; padding: 28px 40px 20px; }
.header h1 { font-size: 20px; font-weight: 700; letter-spacing: .5px; margin-bottom: 4px; }
.header .sub { font-size: 13px; color: #A8B8D0; }
.header .fecha { font-size: 11px; color: #7A90B0; margin-top: 3px; }

/* Tabs */
.tab-nav { display: flex; background: #131f35; overflow-x: auto; }
.tab-btn { padding: 11px 20px; cursor: pointer; color: #8A9BB5; font-size: 12.5px;
           font-weight: 600; border: none; background: none; border-bottom: 3px solid transparent;
           white-space: nowrap; letter-spacing: .2px; }
.tab-btn:hover { color: #D0DCF0; }
.tab-btn.active { color: #fff; border-bottom-color: #4A90D9; background: #1A2B4A; }
.tab-content { display: none; padding: 28px 36px; }
.tab-content.active { display: block; }

/* Secciones */
.section { margin-bottom: 22px; border-radius: 7px; border: 1px solid #E5E8EE; overflow: hidden; }
.section-title { background: #1A2B4A; color: #fff; padding: 9px 16px; font-size: 13.5px;
                 font-weight: 700; letter-spacing: .3px; }
.section-body { padding: 12px 16px; }
.bloque-label { font-size: 11.5px; font-weight: 700; text-transform: uppercase;
                letter-spacing: .8px; margin: 8px 0 5px; }
.label-positivo { color: #1A7A3C; }
.label-mejorar  { color: #C0392B; }
ul.bullets { margin-left: 16px; }
ul.bullets li { font-size: 13px; margin-bottom: 4px; line-height: 1.5; }
.sin-datos { color: #888; font-style: italic; font-size: 13px; }

/* Tabla ICC */
.tbl-icc { width: 100%; border-collapse: collapse; margin-top: 6px; font-size: 12px; }
.tbl-icc th { background: #1A2B4A; color: #fff; padding: 7px 9px; text-align: left; font-weight: 600; }
.tbl-icc td { padding: 6px 9px; border-bottom: 1px solid #EEE; vertical-align: top; }
.tbl-icc tr:last-child td { border-bottom: none; }
.tbl-icc tr:nth-child(even) td { background: #F9FAFB; }

/* Tabla tareas */
.tbl-tareas { width: 100%; border-collapse: collapse; font-size: 12px; margin-top: 6px; }
.tbl-tareas th { background: #1A2B4A; color: #fff; padding: 8px 10px; text-align: left; }
.tbl-tareas td { padding: 7px 10px; border-bottom: 1px solid #EEE; vertical-align: top; }
.tbl-tareas tr:nth-child(even) td { background: #F5F5F5; }

/* Badges */
.badge { display: inline-block; border-radius: 4px; padding: 2px 7px; font-size: 11px; font-weight: 700; }
.badge-pos { background: #D4EDDA; color: #1A7A3C; }
.badge-mej { background: #FADBD8; color: #C0392B; }
.badge-top { background: #FFD700; color: #333; }
.badge-mod { background: #E8ECF5; color: #1A2B4A; font-size: 11px; }

/* Tabla evolución */
.tbl-evo { width: 100%; border-collapse: collapse; font-size: 12.5px; margin-top: 10px; }
.tbl-evo th { background: #1A2B4A; color: #fff; padding: 8px 12px; text-align: center; font-weight: 600; }
.tbl-evo th:first-child { text-align: left; }
.tbl-evo td { padding: 7px 12px; border-bottom: 1px solid #EEE; text-align: center; }
.tbl-evo td:first-child { text-align: left; font-weight: 500; }
.tbl-evo tr:nth-child(even) td { background: #F9FAFB; }
.evo-up   { background: #D4EDDA; color: #1A7A3C; font-weight: 600; }
.evo-down { background: #FADBD8; color: #C0392B; font-weight: 600; }
.evo-null { color: #aaa; }
.evo-obj td { background: #1A2B4A !important; color: #FFD700; font-weight: 700; }
.evo-obj td:first-child { color: #fff; }

/* Tabla objetivos */
.tbl-obj { width: 100%; border-collapse: collapse; font-size: 12.5px; margin-top: 10px; }
.tbl-obj th { background: #1A2B4A; color: #fff; padding: 8px 12px; text-align: left; font-weight: 600; }
.tbl-obj td { padding: 7px 12px; border-bottom: 1px solid #EEE; vertical-align: top; }
.tbl-obj tr:nth-child(even) td { background: #F9FAFB; }
.obj-pos { color: #1A7A3C; font-weight: 600; }
.obj-mej { color: #C0392B; font-weight: 600; }
.obj-note { background: #FFF8E6; border-left: 3px solid #F5A623;
            padding: 10px 14px; margin: 10px 0; font-size: 12.5px; border-radius: 4px; }
"""

JS = """
function showTab(tabId, btn) {
  document.querySelectorAll('.tab-content').forEach(function(el){ el.classList.remove('active'); });
  document.querySelectorAll('.tab-btn').forEach(function(el){ el.classList.remove('active'); });
  document.getElementById(tabId).classList.add('active');
  btn.classList.add('active');
}
"""

def estado_badge(estado):
    if estado == "positivo":
        return '<span class="badge badge-pos">Positivo</span>'
    if estado == "mejorar":
        return '<span class="badge badge-mej">Mejorar</span>'
    return ""


def render_bullets_html(bullets, bloque_class):
    rows_icc = []
    rows_text = []
    icc_mod_header = None

    for txt in bullets:
        parsed = parse_icc_bullet(txt)
        if parsed:
            rows_icc.append(parsed)
        else:
            mod = parse_icc_module_header(txt)
            if mod:
                icc_mod_header = mod
            else:
                rows_text.append(txt)

    html = ""
    if icc_mod_header:
        mod, pos, tot, pts = icc_mod_header
        pts_str = f" | Puntuación: {pts}" if pts else ""
        html += f'<p class="badge badge-mod" style="margin-bottom:8px;">ICC Módulo {mod} — Posición {pos} de {tot}{pts_str}</p>\n'

    if rows_icc:
        html += '<table class="tbl-icc"><thead><tr>'
        html += '<th>Indicador</th><th>Valor</th><th>Media red</th><th>Top 20</th><th>Estado</th>'
        html += '</tr></thead><tbody>\n'
        for ind, val, media, top20, comt in rows_icc:
            top20_cell = f'<span class="badge badge-top">{top20}</span>' if top20 else ""
            est_badge = estado_badge(bloque_class)
            html += f"<tr><td>{ind}</td><td>{val}</td><td>{media}</td><td>{top20_cell}</td><td>{est_badge}</td></tr>\n"
        html += "</tbody></table>\n"

    if rows_text:
        html += '<ul class="bullets">\n'
        for t in rows_text:
            html += f"<li>{t}</li>\n"
        html += "</ul>\n"

    return html


def render_mes_html_content(meta, sections, tabla_tareas):
    """Devuelve el HTML del cuerpo de un mes (sin tags de página ni cabecera)."""
    lines = []
    for sec in sections:
        nombre   = sec["nombre"]
        positivos = sec.get("positivo", [])
        mejoras   = sec.get("mejorar", [])

        lines.append('<div class="section">')
        lines.append(f'<div class="section-title">{nombre}</div>')
        lines.append('<div class="section-body">')

        if nombre == "Tareas pendientes" and tabla_tareas:
            lines.append('<table class="tbl-tareas"><thead><tr>')
            for col in tabla_tareas[0]:
                lines.append(f"<th>{col}</th>")
            lines.append("</tr></thead><tbody>")
            for fila in tabla_tareas[1:]:
                lines.append("<tr>")
                for celda in fila:
                    lines.append(f"<td>{celda}</td>")
                lines.append("</tr>")
            lines.append("</tbody></table>")
        else:
            if not positivos and not mejoras:
                lines.append('<p class="sin-datos">Sin datos registrados</p>')
            else:
                if positivos:
                    lines.append('<p class="bloque-label label-positivo">Puntos positivos</p>')
                    lines.append(render_bullets_html(positivos, "positivo"))
                if mejoras:
                    lines.append('<p class="bloque-label label-mejorar">Puntos a mejorar</p>')
                    lines.append(render_bullets_html(mejoras, "mejorar"))

        lines.append("</div></div>")
    return "\n".join(lines)


def render_evolucion_html(all_data, tipo, key):
    """Tabla comparativa por bloques/secciones del informe."""
    metrics_def = evo_metrics_for_tipo(tipo)
    periodos = [d[0] for d in all_data]
    periodos_label_list = [periodo_label(p) for p in periodos]

    data = {}
    hib_map = {}
    for name, _, hib in metrics_def:
        data[name] = {}
        hib_map[name] = hib
    for periodo, meta, sections, tareas, metricas in all_data:
        for name in data:
            data[name][periodo] = metricas.get(name)

    obj_dict = dict(EVO_OBJ.get(tipo, {}))
    if key in OBJ_PROD_ESPECIFICO:
        obj_dict["Productividad (%)"] = f"{OBJ_PROD_ESPECIFICO[key]}%"

    blocks = EVO_BLOCKS.get(tipo, [])
    if not blocks:
        return '<p class="sin-datos">Sin métricas disponibles para mostrar.</p>'

    lines = []
    any_block_rendered = False

    for block_name, metric_names in blocks:
        filas = [(name, hib_map.get(name, True)) for name in metric_names
                 if name in data and any(v is not None for v in data[name].values())]
        if not filas:
            continue

        any_block_rendered = True
        lines.append('<div class="section" style="margin-bottom:16px;">')
        lines.append(f'<div class="section-title">{block_name}</div>')
        lines.append('<div class="section-body">')
        lines.append('<table class="tbl-evo"><thead><tr>')
        lines.append('<th>Métrica</th>')
        for pl in periodos_label_list:
            lines.append(f'<th>{pl}</th>')
        lines.append('<th>Objetivo</th>')
        lines.append('</tr></thead><tbody>')

        for name, higher_better in filas:
            lines.append('<tr>')
            lines.append(f'<td>{name}</td>')
            values = [data[name].get(p) for p in periodos]
            for i, val in enumerate(values):
                if val is None:
                    lines.append('<td class="evo-null">—</td>')
                    continue
                cls = ""
                if i > 0:
                    prev = values[i - 1]
                    if prev is not None and isinstance(val, float) and isinstance(prev, float):
                        if val > prev:
                            cls = "evo-up" if higher_better else "evo-down"
                        elif val < prev:
                            cls = "evo-down" if higher_better else "evo-up"
                display = str(val).replace(".", ",") if isinstance(val, float) else str(val)
                lines.append(f'<td class="{cls}">{display}</td>')
            obj_val = obj_dict.get(name, "—")
            lines.append(f'<td style="background:#1A2B4A;color:#FFD700;font-weight:700;">{obj_val}</td>')
            lines.append('</tr>')

        lines.append('</tbody></table>')
        lines.append('</div></div>')

    if not any_block_rendered:
        return '<p class="sin-datos">Sin métricas disponibles para mostrar.</p>'

    return "\n".join(lines)


def render_objetivos_html(tipo, key):
    """Tarjeta de referencia de umbrales por tipo."""
    rows = OBJETIVOS.get(tipo, [])
    if not rows:
        return '<p class="sin-datos">Sin objetivos definidos para este tipo.</p>'

    lines = []

    # Nota de productividad específica
    if tipo in ("jt_mecanica", "jt_chapa") and key in OBJ_PROD_ESPECIFICO:
        obj_prod = OBJ_PROD_ESPECIFICO[key]
        lines.append(f'<div class="obj-note">Objetivo de productividad específico para esta instalación: <strong>{obj_prod}%</strong></div>')

    lines.append('<table class="tbl-obj"><thead><tr>')
    lines.append('<th>Métrica</th><th>Referencia</th>')
    lines.append('<th class="obj-pos">Positivo si...</th>')
    lines.append('<th class="obj-mej">Mejorar si...</th>')
    lines.append('</tr></thead><tbody>')
    for metrica, referencia, positivo, mejorar in rows:
        lines.append('<tr>')
        lines.append(f'<td><strong>{metrica}</strong></td>')
        lines.append(f'<td>{referencia}</td>')
        lines.append(f'<td class="obj-pos">{positivo}</td>')
        lines.append(f'<td class="obj-mej">{mejorar}</td>')
        lines.append('</tr>')
    lines.append('</tbody></table>')
    return "\n".join(lines)


def gen_html_multitab(key, all_data, tipo, base):
    """Genera el HTML multi-pestaña para una persona."""
    # Meta principal: usar el del mes más reciente
    last = all_data[-1]
    meta = last[1]

    out_path = os.path.join(base, f"informe_seguimiento_{key}.html")

    lines = [
        "<!DOCTYPE html>",
        '<html lang="es"><head>',
        '<meta charset="UTF-8">',
        f'<title>{meta["titulo"]}</title>',
        f'<style>{CSS}</style>',
        "</head><body>",
        f'<script>{JS}</script>',
        '<div class="page">',
        '<div class="header">',
        f'<h1>{meta["titulo"]}</h1>',
        f'<div class="sub">{meta["subtitulo"]}</div>',
        f'<div class="fecha">Generado: {datetime.date.today().strftime("%d/%m/%Y")}</div>',
        '</div>',  # header
    ]

    # Tab nav
    lines.append('<nav class="tab-nav">')
    for idx, (periodo, meta_p, sections, tareas, metricas) in enumerate(all_data):
        pl = periodo_label(periodo)
        active = "active" if idx == len(all_data) - 1 else ""
        tid = f"tab-{periodo}"
        lines.append(f'<button class="tab-btn {active}" onclick="showTab(\'{tid}\', this)">{pl}</button>')
    lines.append('<button class="tab-btn" onclick="showTab(\'tab-evo\', this)">Evolución</button>')
    lines.append('</nav>')

    # Tab contents — meses
    for idx, (periodo, meta_p, sections, tareas, metricas) in enumerate(all_data):
        tid = f"tab-{periodo}"
        active = "active" if idx == len(all_data) - 1 else ""
        lines.append(f'<div id="{tid}" class="tab-content {active}">')
        lines.append(render_mes_html_content(meta_p, sections, tareas))
        lines.append('</div>')

    # Tab evolución
    lines.append('<div id="tab-evo" class="tab-content">')
    lines.append('<div class="section"><div class="section-title">Evolución por periodo</div>')
    lines.append('<div class="section-body">')
    lines.append(render_evolucion_html(all_data, tipo, key))
    lines.append('</div></div></div>')

    lines.append('</div></body></html>')

    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    return out_path


# ══════════════════════════════════════════════════════════════════════════════
# EXCEL — HELPERS
# ══════════════════════════════════════════════════════════════════════════════

FILL_HEADER  = PatternFill("solid", fgColor="1A2B4A")
FILL_SEC_HDR = PatternFill("solid", fgColor="1A2B4A")
FILL_TBL_HDR = PatternFill("solid", fgColor="2C3E6B")
FILL_POS     = PatternFill("solid", fgColor="D4EDDA")
FILL_MEJ     = PatternFill("solid", fgColor="FADBD8")
FILL_TOP     = PatternFill("solid", fgColor="FFD700")
FILL_ALT     = PatternFill("solid", fgColor="F9FAFB")
FILL_EVO_UP  = PatternFill("solid", fgColor="D4EDDA")
FILL_EVO_DW  = PatternFill("solid", fgColor="FADBD8")
FILL_OBJ_POS = PatternFill("solid", fgColor="EAF7EE")
FILL_OBJ_MEJ = PatternFill("solid", fgColor="FDECEA")
FILL_NOTE    = PatternFill("solid", fgColor="FFF8E6")

FONT_WHITE_BOLD = Font(name="Calibri", bold=True, color="FFFFFF", size=11)
FONT_WHITE_SM   = Font(name="Calibri", color="FFFFFF", size=10)
FONT_DARK_BOLD  = Font(name="Calibri", bold=True, color="1A2B4A", size=10)
FONT_DARK       = Font(name="Calibri", color="222222", size=10)
FONT_POS        = Font(name="Calibri", bold=True, color="1A7A3C", size=10)
FONT_MEJ        = Font(name="Calibri", bold=True, color="C0392B", size=10)
FONT_ITALIC     = Font(name="Calibri", italic=True, color="888888", size=10)

THIN  = Side(border_style="thin",   color="DDDDDD")
BORDER_THIN = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)

NCOLS = 6

def xl_write(ws, row, col, value, fill=None, font=None, align=None, border=None, wrap=False):
    c = ws.cell(row=row, column=col, value=value)
    if fill:   c.fill = fill
    if font:   c.font = font
    if align:  c.alignment = align
    if border: c.border = border
    if wrap:   c.alignment = Alignment(wrap_text=True, vertical="top")
    return c

def xl_merge_write(ws, row, col1, col2, value, fill=None, font=None):
    ws.merge_cells(start_row=row, start_column=col1, end_row=row, end_column=col2)
    c = ws.cell(row=row, column=col1, value=value)
    if fill: c.fill = fill
    if font: c.font = font
    c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    return c


def set_col_widths(ws):
    ws.column_dimensions["A"].width = 4
    ws.column_dimensions["B"].width = 42
    ws.column_dimensions["C"].width = 16
    ws.column_dimensions["D"].width = 16
    ws.column_dimensions["E"].width = 14
    ws.column_dimensions["F"].width = 14


def render_mes_sheet(ws, meta, sections, tabla_tareas):
    """Escribe el contenido de un mes en una hoja Excel existente."""
    set_col_widths(ws)
    row = 1

    ws.row_dimensions[row].height = 30
    xl_merge_write(ws, row, 1, NCOLS, meta["titulo"], fill=FILL_HEADER,
                   font=Font(name="Calibri", bold=True, color="FFFFFF", size=14))
    row += 1
    ws.row_dimensions[row].height = 18
    xl_merge_write(ws, row, 1, NCOLS, meta["subtitulo"], fill=FILL_HEADER,
                   font=Font(name="Calibri", color="A8B8D0", size=11))
    row += 1
    ws.row_dimensions[row].height = 16
    xl_merge_write(ws, row, 1, NCOLS, meta["fecha"], fill=FILL_HEADER,
                   font=Font(name="Calibri", color="7A90B0", size=10))
    row += 2  # espacio

    sec_num = 0
    for sec in sections:
        nombre   = sec["nombre"]
        positivos = sec.get("positivo", [])
        mejoras   = sec.get("mejorar", [])
        sec_num += 1

        ws.row_dimensions[row].height = 20
        xl_merge_write(ws, row, 1, NCOLS, f"{sec_num}. {nombre}",
                       fill=FILL_SEC_HDR,
                       font=Font(name="Calibri", bold=True, color="FFFFFF", size=11))
        row += 1

        if nombre == "Tareas pendientes" and tabla_tareas:
            ws.row_dimensions[row].height = 18
            for ci, col_name in enumerate(tabla_tareas[0]):
                xl_write(ws, row, ci + 1, col_name, fill=FILL_TBL_HDR,
                         font=FONT_WHITE_BOLD, border=BORDER_THIN,
                         align=Alignment(horizontal="center", vertical="center"))
            row += 1
            for fi, fila in enumerate(tabla_tareas[1:]):
                ws.row_dimensions[row].height = 30
                fill_row = FILL_ALT if fi % 2 == 1 else None
                for ci, celda in enumerate(fila):
                    xl_write(ws, row, ci + 1, celda, fill=fill_row,
                             font=FONT_DARK, border=BORDER_THIN, wrap=True)
                row += 1
        elif not positivos and not mejoras:
            ws.row_dimensions[row].height = 16
            xl_merge_write(ws, row, 1, NCOLS, "Sin datos registrados", font=FONT_ITALIC)
            row += 1
        else:
            for bloque_key, bloque_name, bloque_font, bloque_fill in [
                ("positivo", "Puntos positivos", FONT_POS, FILL_POS),
                ("mejorar",  "Puntos a mejorar", FONT_MEJ, FILL_MEJ),
            ]:
                bullets = sec.get(bloque_key, [])
                if not bullets:
                    continue

                ws.row_dimensions[row].height = 16
                xl_merge_write(ws, row, 1, NCOLS, bloque_name.upper(), font=bloque_font)
                row += 1

                rows_icc = []
                icc_header = None
                text_bullets = []

                for txt in bullets:
                    parsed = parse_icc_bullet(txt)
                    if parsed:
                        rows_icc.append(parsed)
                    else:
                        mod = parse_icc_module_header(txt)
                        if mod:
                            icc_header = mod
                        else:
                            text_bullets.append(txt)

                if icc_header:
                    mod, pos, tot, pts = icc_header
                    pts_str = f" | Puntuación: {pts}" if pts else ""
                    ws.row_dimensions[row].height = 16
                    xl_merge_write(ws, row, 1, NCOLS,
                                   f"ICC Módulo {mod} — Posición {pos} de {tot}{pts_str}",
                                   fill=PatternFill("solid", fgColor="E8ECF5"),
                                   font=Font(name="Calibri", bold=True, color="1A2B4A", size=10))
                    row += 1

                if rows_icc:
                    ws.row_dimensions[row].height = 16
                    for ci, hdr in enumerate(["#", "Indicador", "Valor", "Media red", "Top 20", "Estado"], 1):
                        xl_write(ws, row, ci, hdr, fill=FILL_TBL_HDR,
                                 font=FONT_WHITE_BOLD, border=BORDER_THIN,
                                 align=Alignment(horizontal="center", vertical="center"))
                    row += 1
                    for ri, (ind, val, media, top20, comt) in enumerate(rows_icc):
                        ws.row_dimensions[row].height = 28
                        fill_row = FILL_ALT if ri % 2 == 1 else None
                        xl_write(ws, row, 1, ri + 1, fill=fill_row, font=FONT_DARK, border=BORDER_THIN,
                                 align=Alignment(horizontal="center", vertical="top"))
                        xl_write(ws, row, 2, ind, fill=fill_row, font=FONT_DARK, border=BORDER_THIN, wrap=True)
                        xl_write(ws, row, 3, val, fill=fill_row, font=FONT_DARK, border=BORDER_THIN, wrap=True)
                        xl_write(ws, row, 4, media, fill=fill_row, font=FONT_DARK, border=BORDER_THIN, wrap=True)
                        top20_val = top20 if top20 else ""
                        xl_write(ws, row, 5, top20_val,
                                 fill=FILL_TOP if top20 else fill_row,
                                 font=Font(name="Calibri", bold=True, color="333333", size=10) if top20 else FONT_DARK,
                                 border=BORDER_THIN,
                                 align=Alignment(horizontal="center", vertical="top"))
                        lbl = "Positivo" if bloque_key == "positivo" else "Mejorar"
                        fll = FILL_POS if bloque_key == "positivo" else FILL_MEJ
                        fnt = FONT_POS if bloque_key == "positivo" else FONT_MEJ
                        xl_write(ws, row, 6, lbl, fill=fll, font=fnt, border=BORDER_THIN,
                                 align=Alignment(horizontal="center", vertical="top"))
                        row += 1

                if text_bullets:
                    for ti, txt in enumerate(text_bullets):
                        ws.row_dimensions[row].height = 28
                        fill_row = FILL_ALT if ti % 2 == 1 else None
                        xl_write(ws, row, 1, "•", fill=fill_row, font=FONT_DARK,
                                 align=Alignment(horizontal="center", vertical="top"))
                        c = ws.cell(row=row, column=2, value=txt)
                        c.fill = fill_row or PatternFill()
                        c.font = FONT_DARK
                        c.alignment = Alignment(wrap_text=True, vertical="top")
                        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=NCOLS)
                        row += 1

        row += 1  # espacio entre secciones


def render_evolucion_sheet(ws, all_data, tipo, key):
    """Escribe la hoja de evolución por bloques/secciones del informe."""
    metrics_def = evo_metrics_for_tipo(tipo)
    periodos = [d[0] for d in all_data]
    periodos_label_list = [periodo_label(p) for p in periodos]
    n_per = len(periodos)
    n_cols = 1 + n_per + 1  # métrica + periodos + objetivo

    ws.column_dimensions["A"].width = 22
    for i in range(n_per):
        ws.column_dimensions[get_column_letter(i + 2)].width = 14
    ws.column_dimensions[get_column_letter(n_per + 2)].width = 14

    data = {}
    hib_map = {}
    for name, _, hib in metrics_def:
        data[name] = {}
        hib_map[name] = hib
    for periodo, meta, sections, tareas, metricas in all_data:
        for name in data:
            data[name][periodo] = metricas.get(name)

    obj_dict = dict(EVO_OBJ.get(tipo, {}))
    if key in OBJ_PROD_ESPECIFICO:
        obj_dict["Productividad (%)"] = f"{OBJ_PROD_ESPECIFICO[key]}%"

    FILL_OBJ_CELL  = PatternFill("solid", fgColor="1A2B4A")
    FONT_OBJ_CELL  = Font(name="Calibri", bold=True, color="FFD700", size=10)
    FILL_BLOCK_HDR = PatternFill("solid", fgColor="1A2B4A")
    FONT_BLOCK_HDR = Font(name="Calibri", bold=True, color="FFFFFF", size=11)

    row = 1
    ws.row_dimensions[row].height = 20
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=n_cols)
    c = ws.cell(row=row, column=1, value="Evolución por periodo")
    c.fill = FILL_HEADER
    c.font = Font(name="Calibri", bold=True, color="FFFFFF", size=13)
    c.alignment = Alignment(horizontal="left", vertical="center")
    row += 2

    blocks = EVO_BLOCKS.get(tipo, [])

    for block_name, metric_names in blocks:
        filas = [(name, hib_map.get(name, True)) for name in metric_names
                 if name in data and any(v is not None for v in data[name].values())]
        if not filas:
            continue

        # Cabecera de bloque (azul oscuro, texto blanco)
        ws.row_dimensions[row].height = 20
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=n_cols)
        c = ws.cell(row=row, column=1, value=block_name)
        c.fill = FILL_BLOCK_HDR
        c.font = FONT_BLOCK_HDR
        c.alignment = Alignment(horizontal="left", vertical="center")
        row += 1

        # Cabecera de columnas
        ws.row_dimensions[row].height = 18
        xl_write(ws, row, 1, "Métrica", fill=FILL_TBL_HDR, font=FONT_WHITE_BOLD, border=BORDER_THIN,
                 align=Alignment(horizontal="left", vertical="center"))
        for i, pl in enumerate(periodos_label_list):
            xl_write(ws, row, i + 2, pl, fill=FILL_TBL_HDR, font=FONT_WHITE_BOLD, border=BORDER_THIN,
                     align=Alignment(horizontal="center", vertical="center"))
        xl_write(ws, row, n_per + 2, "Objetivo", fill=FILL_OBJ_CELL, font=FONT_OBJ_CELL,
                 border=BORDER_THIN, align=Alignment(horizontal="center", vertical="center"))
        row += 1

        for ri, (name, higher_better) in enumerate(filas):
            ws.row_dimensions[row].height = 22
            fill_base = FILL_ALT if ri % 2 == 1 else None
            xl_write(ws, row, 1, name, fill=fill_base, font=FONT_DARK_BOLD, border=BORDER_THIN)
            values = [data[name].get(p) for p in periodos]
            for i, val in enumerate(values):
                if val is None:
                    xl_write(ws, row, i + 2, "—", fill=fill_base,
                             font=FONT_ITALIC, border=BORDER_THIN,
                             align=Alignment(horizontal="center"))
                    continue
                fill_cell = fill_base
                font_cell = FONT_DARK
                if i > 0:
                    prev = values[i - 1]
                    if prev is not None and isinstance(val, float) and isinstance(prev, float):
                        if val > prev:
                            fill_cell = FILL_EVO_UP if higher_better else FILL_EVO_DW
                            font_cell = FONT_POS if higher_better else FONT_MEJ
                        elif val < prev:
                            fill_cell = FILL_EVO_DW if higher_better else FILL_EVO_UP
                            font_cell = FONT_MEJ if higher_better else FONT_POS
                xl_write(ws, row, i + 2, val, fill=fill_cell, font=font_cell, border=BORDER_THIN,
                         align=Alignment(horizontal="center"))
            obj_val = obj_dict.get(name, "—")
            xl_write(ws, row, n_per + 2, obj_val, fill=FILL_OBJ_CELL, font=FONT_OBJ_CELL,
                     border=BORDER_THIN, align=Alignment(horizontal="center"))
            row += 1

        row += 1  # espacio entre bloques


def render_objetivos_sheet(ws, tipo, key):
    """Escribe la hoja de objetivos."""
    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 28
    ws.column_dimensions["C"].width = 28
    ws.column_dimensions["D"].width = 32

    rows = OBJETIVOS.get(tipo, [])
    row = 1
    ws.row_dimensions[row].height = 20
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
    c = ws.cell(row=row, column=1, value="Objetivos y umbrales de referencia")
    c.fill = FILL_HEADER
    c.font = Font(name="Calibri", bold=True, color="FFFFFF", size=13)
    c.alignment = Alignment(horizontal="left", vertical="center")
    row += 2

    # Nota productividad
    if tipo in ("jt_mecanica", "jt_chapa") and key in OBJ_PROD_ESPECIFICO:
        obj_prod = OBJ_PROD_ESPECIFICO[key]
        ws.row_dimensions[row].height = 24
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        c = ws.cell(row=row, column=1,
                    value=f"Objetivo de productividad específico para esta instalación: {obj_prod}%")
        c.fill = FILL_NOTE
        c.font = Font(name="Calibri", bold=True, color="8B6914", size=10)
        c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        row += 2

    # Cabecera tabla
    ws.row_dimensions[row].height = 18
    for ci, hdr in enumerate(["Métrica", "Referencia", "Positivo si...", "Mejorar si..."], 1):
        xl_write(ws, row, ci, hdr, fill=FILL_TBL_HDR, font=FONT_WHITE_BOLD, border=BORDER_THIN,
                 align=Alignment(horizontal="left", vertical="center"))
    row += 1

    for ri, (metrica, referencia, positivo, mejorar) in enumerate(rows):
        ws.row_dimensions[row].height = 24
        fill_base = FILL_ALT if ri % 2 == 1 else None
        xl_write(ws, row, 1, metrica, fill=fill_base, font=FONT_DARK_BOLD, border=BORDER_THIN, wrap=True)
        xl_write(ws, row, 2, referencia, fill=fill_base, font=FONT_DARK, border=BORDER_THIN, wrap=True)
        xl_write(ws, row, 3, positivo, fill=FILL_OBJ_POS, font=FONT_POS, border=BORDER_THIN, wrap=True)
        xl_write(ws, row, 4, mejorar, fill=FILL_OBJ_MEJ, font=FONT_MEJ, border=BORDER_THIN, wrap=True)
        row += 1


def gen_excel_multitab(key, all_data, tipo, base):
    """Genera el Excel multi-hoja para una persona."""
    out_path = os.path.join(base, f"informe_seguimiento_{key}.xlsx")
    wb = Workbook()

    for idx, (periodo, meta, sections, tareas, metricas) in enumerate(all_data):
        title = periodo_label(periodo)
        ws = wb.create_sheet(title=title)
        if idx == 0:
            # Eliminar la hoja vacía inicial
            del wb[wb.sheetnames[0]]
            ws = wb.create_sheet(title=title)
        render_mes_sheet(ws, meta, sections, tareas)

    ws_evo = wb.create_sheet(title="Evolución")
    render_evolucion_sheet(ws_evo, all_data, tipo, key)

    wb.save(out_path)
    return out_path


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    base = os.path.dirname(os.path.abspath(__file__))
    groups = group_files(base)

    print(f"Personas encontradas: {len(groups)}\n")
    ok = 0
    errors = []

    for key, periodos in sorted(groups.items()):
        all_data = []
        for periodo, path in periodos:
            try:
                meta, sections, tareas = parse_word(path)
                metricas = extract_evolution_metrics(sections, detect_tipo(key))
                all_data.append((periodo, meta, sections, tareas, metricas))
            except Exception as e:
                errors.append((key, periodo, str(e)))
                continue

        if not all_data:
            continue

        tipo = detect_tipo(key)
        try:
            gen_html_multitab(key, all_data, tipo, base)
            gen_excel_multitab(key, all_data, tipo, base)
            periodos_str = " + ".join(periodo_label(p) for p, _ in periodos)
            print(f"  OK  {key}  [{periodos_str}]")
            ok += 1
        except Exception as e:
            print(f"  ERR {key}: {e}")
            errors.append((key, "output", str(e)))

    print(f"\nGenerados: {ok}/{len(groups)}")
    if errors:
        print("\nErrores:")
        for parts in errors:
            print(f"  {' | '.join(parts)}")


if __name__ == "__main__":
    main()
